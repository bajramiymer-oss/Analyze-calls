
import os, pathlib, requests
from typing import List
from docx import Document
from .config import load_openai_key

def transcribe_audio_files(audio_paths: List[pathlib.Path], out_dir: pathlib.Path, mode: str = "per-file") -> dict:
    """
    Transkripton me Whisper-1. Kthen dict:
    {
      "txt_paths": [ ... ],
      "merged_docx": pathlib.Path | None
    }
    """
    out_dir.mkdir(parents=True, exist_ok=True)
    api_key = load_openai_key()
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY mungon")
    headers = {"Authorization": f"Bearer {api_key}"}

    txt_paths = []
    for p in audio_paths:
        with open(p, "rb") as f:
            files = {"file": (p.name, f, "application/octet-stream")}
            data = {"model": "whisper-1", "response_format": "text"}
            r = requests.post("https://api.openai.com/v1/audio/transcriptions", headers=headers, data=data, files=files, timeout=300)
        if r.status_code != 200:
            raise RuntimeError(f"Whisper error {r.status_code}: {r.text[:200]}")
        text = r.text
        out_txt = out_dir / (p.stem + ".txt")
        out_txt.write_text(text, encoding="utf-8")
        txt_paths.append(out_txt)

    merged_docx = None
    if mode == "merged":
        doc = Document()
        for t in txt_paths:
            doc.add_heading(t.name, level=2)
            doc.add_paragraph(t.read_text(encoding="utf-8"))
            doc.add_page_break()
        merged_docx = out_dir / "merged_transcripts.docx"
        doc.save(str(merged_docx))

    return {"txt_paths": txt_paths, "merged_docx": merged_docx}
