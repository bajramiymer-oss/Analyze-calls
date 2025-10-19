"""
core/campaign_manager.py

PURPOSE:
    Campaign Context Manager - Full CRUD operations for campaigns
    and document management (PDF, DOCX, TXT).

RESPONSIBILITIES:
    - Create, read, update, delete campaigns
    - Upload and extract text from documents
    - Manage campaign metadata and limits
    - Provide campaign context for LLM analysis

KEY FUNCTIONS:
    - create_campaign() - Create new campaign with hints
    - add_document_to_campaign() - Upload & extract document text
    - get_campaign_hints() - Get all context for a campaign
    - delete_campaign() - Remove campaign & all documents

LIMITS:
    - Max 3 documents per campaign
    - Max 5MB per document
    - Max 50 pages total across all documents

STORAGE:
    - Metadata: config/campaign_contexts.json
    - Documents: assets/campaigns/{campaign_id}/documents/
    - Extracted text: assets/campaigns/{campaign_id}/extracted_text.txt

DEPENDENCIES:
    - PyPDF2 (for PDF extraction)
    - python-docx (for Word extraction)

Author: Protrade AI
Last Updated: 2025-10-13
"""

import json
import pathlib
import shutil
from datetime import datetime
from typing import Dict, List, Optional
import hashlib

# Imports për procesim dokumentesh
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# ============== CONSTANTS ==============
CAMPAIGN_CONFIG_PATH = pathlib.Path("config/campaign_contexts.json")
CAMPAIGN_ASSETS_DIR = pathlib.Path("assets/campaigns")

MAX_DOCUMENTS_PER_CAMPAIGN = 3
MAX_FILE_SIZE_MB = 5
MAX_TOTAL_PAGES = 50


# ============== CONFIG HELPERS ==============
def ensure_campaign_config():
    """Krijon file-in e konfigurimit nëse nuk ekziston"""
    CAMPAIGN_CONFIG_PATH.parent.mkdir(parents=True, exist_ok=True)
    if not CAMPAIGN_CONFIG_PATH.exists():
        default_config = {
            "campaigns": [],
            "settings": {
                "max_documents_per_campaign": MAX_DOCUMENTS_PER_CAMPAIGN,
                "max_file_size_mb": MAX_FILE_SIZE_MB,
                "max_total_pages": MAX_TOTAL_PAGES
            }
        }
        with open(CAMPAIGN_CONFIG_PATH, 'w', encoding='utf-8') as f:
            json.dump(default_config, f, indent=2, ensure_ascii=False)


def load_campaign_config() -> dict:
    """Ngarkon konfigurimin e fushatave"""
    ensure_campaign_config()
    with open(CAMPAIGN_CONFIG_PATH, 'r', encoding='utf-8') as f:
        return json.load(f)


def save_campaign_config(config: dict):
    """Ruaj konfigurimin e fushatave"""
    ensure_campaign_config()
    with open(CAMPAIGN_CONFIG_PATH, 'w', encoding='utf-8') as f:
        json.dump(config, f, indent=2, ensure_ascii=False)


# ============== CAMPAIGN OPERATIONS ==============
def get_all_campaigns() -> List[dict]:
    """Merr të gjitha fushatat"""
    config = load_campaign_config()
    return config.get("campaigns", [])


def get_campaign_by_id(campaign_id: str) -> Optional[dict]:
    """Merr një fushatë sipas ID-së"""
    campaigns = get_all_campaigns()
    for c in campaigns:
        if c.get("id") == campaign_id:
            return c
    return None


def get_campaign_by_name(name: str) -> Optional[dict]:
    """Merr një fushatë sipas emrit"""
    campaigns = get_all_campaigns()
    for c in campaigns:
        if c.get("name") == name:
            return c
    return None


def generate_campaign_id(name: str) -> str:
    """Gjeneron një ID unik për fushatën"""
    # Krijon një ID të shkurtër bazuar në emër + timestamp
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    name_slug = name.lower().replace(" ", "_").replace("-", "_")
    # Merr vetëm karaktere alfanumerike
    name_slug = ''.join(c for c in name_slug if c.isalnum() or c == '_')
    # Limit në 30 karaktere
    name_slug = name_slug[:30]
    # Hash i shkurtër për unicitet
    hash_part = hashlib.md5(f"{name}{timestamp}".encode()).hexdigest()[:6]
    return f"{name_slug}_{hash_part}"


def create_campaign(
    name: str,
    project_context_hint: str = "",
    summary_hint: str = "",
    bullets_hint: str = ""
) -> dict:
    """
    Krijon një fushatë të re me kontekst dhe metadata.

    Creates a new campaign with unique ID, folder structure, and metadata.
    The campaign ID is auto-generated from name + timestamp + hash.
    Campaign folder is created at assets/campaigns/{campaign_id}/

    Args:
        name: Emri i fushatës (unique, required)
            Example: "Google Reserve UK - Restorante"
        project_context_hint: Konteksti i biznesit/fushatës (optional)
            Example: "Thirrje outbound për restorante në MB..."
        summary_hint: Udhëzime për përmbledhjen LLM (optional)
            Example: "Fokuso në aftësinë e agjentit për të shpjeguar..."
        bullets_hint: Udhëzime për bullet points LLM (optional)
            Example: "Preggi: edukim digjital, qartësi..."

    Returns:
        dict: Campaign object with structure:
            {
                "id": "google_reserve_uk_abc123",
                "name": "Google Reserve UK",
                "created_date": "2025-10-13T10:30:00",
                "last_modified": "2025-10-13T10:30:00",
                "project_context_hint": "...",
                "summary_hint": "...",
                "bullets_hint": "...",
                "documents": [],
                "total_pages": 0,
                "total_size_kb": 0
            }

    Raises:
        ValueError: Nëse ekziston tashmë fushatë me këtë emër

    Example:
        >>> campaign = create_campaign(
        ...     name="Tim Business Albania",
        ...     project_context_hint="Thirrje B2B për aktivizim internet..."
        ... )
        >>> print(campaign["id"])
        tim_business_albania_a1b2c3

        >>> # Upload document
        >>> add_document_to_campaign(
        ...     campaign["id"],
        ...     Path("script.pdf")
        ... )

    Related:
        - update_campaign() - Modify existing campaign
        - delete_campaign() - Remove campaign
        - add_document_to_campaign() - Upload documents
        - get_campaign_by_id() - Retrieve campaign

    Note:
        - Campaign ID është i përhershëm (nuk ndryshon me rename)
        - Folderi krijohet automatikisht te assets/campaigns/
        - Hints mund të lihen të zbrazëta (optional)
    """
    # Kontrollo nëse ekziston fushatë me këtë emër
    if get_campaign_by_name(name):
        raise ValueError(f"Ekziston tashmë një fushatë me emrin: {name}")

    campaign_id = generate_campaign_id(name)
    now = datetime.now().isoformat()

    campaign = {
        "id": campaign_id,
        "name": name,
        "created_date": now,
        "last_modified": now,
        "project_context_hint": project_context_hint.strip(),
        "summary_hint": summary_hint.strip(),
        "bullets_hint": bullets_hint.strip(),
        "documents": [],
        "total_pages": 0,
        "total_size_kb": 0
    }

    # Krijo folderin e fushatës
    campaign_dir = CAMPAIGN_ASSETS_DIR / campaign_id
    campaign_dir.mkdir(parents=True, exist_ok=True)
    (campaign_dir / "documents").mkdir(exist_ok=True)

    # Ruaj në config
    config = load_campaign_config()
    config["campaigns"].append(campaign)
    save_campaign_config(config)

    return campaign


def update_campaign(
    campaign_id: str,
    name: Optional[str] = None,
    project_context_hint: Optional[str] = None,
    summary_hint: Optional[str] = None,
    bullets_hint: Optional[str] = None
) -> dict:
    """
    Përditëson një fushatë ekzistuese

    Args:
        campaign_id: ID e fushatës
        name: Emri i ri (opsional)
        project_context_hint: Konteksti i ri (opsional)
        summary_hint: Udhëzimet e reja (opsional)
        bullets_hint: Udhëzimet e reja (opsional)

    Returns:
        dict: Kampanja e përditësuar
    """
    config = load_campaign_config()
    campaigns = config.get("campaigns", [])

    campaign_idx = None
    for idx, c in enumerate(campaigns):
        if c.get("id") == campaign_id:
            campaign_idx = idx
            break

    if campaign_idx is None:
        raise ValueError(f"Nuk u gjet fushatë me ID: {campaign_id}")

    campaign = campaigns[campaign_idx]

    # Përditëso fushat
    if name is not None:
        # Kontrollo që emri i ri të mos përdoret nga fushatë tjetër
        for c in campaigns:
            if c.get("id") != campaign_id and c.get("name") == name:
                raise ValueError(f"Ekziston tashmë një fushatë me emrin: {name}")
        campaign["name"] = name

    if project_context_hint is not None:
        campaign["project_context_hint"] = project_context_hint.strip()

    if summary_hint is not None:
        campaign["summary_hint"] = summary_hint.strip()

    if bullets_hint is not None:
        campaign["bullets_hint"] = bullets_hint.strip()

    campaign["last_modified"] = datetime.now().isoformat()

    campaigns[campaign_idx] = campaign
    config["campaigns"] = campaigns
    save_campaign_config(config)

    return campaign


def delete_campaign(campaign_id: str):
    """Fshin një fushatë dhe të gjitha dokumentet e saj"""
    config = load_campaign_config()
    campaigns = config.get("campaigns", [])

    # Gjej fushatën
    campaign = None
    for idx, c in enumerate(campaigns):
        if c.get("id") == campaign_id:
            campaign = c
            campaigns.pop(idx)
            break

    if campaign is None:
        raise ValueError(f"Nuk u gjet fushatë me ID: {campaign_id}")

    # Fshi folderin e fushatës
    campaign_dir = CAMPAIGN_ASSETS_DIR / campaign_id
    if campaign_dir.exists():
        shutil.rmtree(campaign_dir)

    # Ruaj konfigurimin e përditësuar
    config["campaigns"] = campaigns
    save_campaign_config(config)


# ============== DOCUMENT PROCESSING ==============
def extract_text_from_pdf(file_path: pathlib.Path) -> tuple[str, int]:
    """
    Ekstrakton tekstin nga një file PDF

    Returns:
        tuple: (teksti, numri i faqeve)
    """
    if not PDF_AVAILABLE:
        raise ImportError("PyPDF2 nuk është instaluar. Run: pip install PyPDF2")

    text_parts = []
    with open(file_path, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        num_pages = len(reader.pages)

        for page_num in range(num_pages):
            page = reader.pages[page_num]
            text = page.extract_text()
            if text:
                text_parts.append(text)

    full_text = "\n\n".join(text_parts)
    return full_text, num_pages


def extract_text_from_docx(file_path: pathlib.Path) -> tuple[str, int]:
    """
    Ekstrakton tekstin nga një file DOCX

    Returns:
        tuple: (teksti, numri i faqeve aproksimative)
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx nuk është instaluar. Run: pip install python-docx")

    doc = Document(str(file_path))
    text_parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    full_text = "\n\n".join(text_parts)

    # Aproksimim i faqeve (300 fjalë për faqe)
    word_count = len(full_text.split())
    estimated_pages = max(1, word_count // 300)

    return full_text, estimated_pages


def extract_text_from_txt(file_path: pathlib.Path) -> tuple[str, int]:
    """
    Merr tekstin nga një file TXT

    Returns:
        tuple: (teksti, numri i faqeve aproksimative)
    """
    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        text = f.read()

    # Aproksimim i faqeve (300 fjalë për faqe)
    word_count = len(text.split())
    estimated_pages = max(1, word_count // 300)

    return text, estimated_pages


def extract_text_from_file(file_path: pathlib.Path) -> tuple[str, int]:
    """
    Ekstrakton tekstin nga një file (PDF, DOCX, ose TXT)

    Returns:
        tuple: (teksti, numri i faqeve)
    """
    suffix = file_path.suffix.lower()

    if suffix == '.pdf':
        return extract_text_from_pdf(file_path)
    elif suffix in ['.docx', '.doc']:
        return extract_text_from_docx(file_path)
    elif suffix == '.txt':
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Format i pambështetur: {suffix}")


def add_document_to_campaign(
    campaign_id: str,
    file_path: pathlib.Path,
    original_filename: Optional[str] = None
) -> dict:
    """
    Shton një dokument në një fushatë dhe ekstrakton tekstin.

    Uploads a document (PDF/DOCX/TXT), extracts text, and appends to
    the campaign's extracted_text.txt file for LLM context.

    Process:
        1. Validates campaign exists and limits not exceeded
        2. Checks file size (max 5MB) and total pages (max 50)
        3. Extracts text based on file type
        4. Copies file to assets/campaigns/{campaign_id}/documents/
        5. Appends extracted text to extracted_text.txt
        6. Updates campaign metadata in JSON

    Args:
        campaign_id: ID e fushatës (e.g., "google_reserve_uk_abc123")
        file_path: Path i file-it që do të ngarkohet
            Supported formats: .pdf, .docx, .txt
        original_filename: Emri origjinal i file-it (optional)
            If None, uses file_path.name

    Returns:
        dict: Updated campaign object with new document in "documents" list

    Raises:
        ValueError: Nëse:
            - Campaign nuk ekziston
            - Më shumë se 3 dokumente (limit exceeded)
            - File më i madh se 5MB
            - Total faqe > 50
            - Format i pambështetur
            - Gabim në ekstraktimin e tekstit

    Example:
        >>> campaign = get_campaign_by_id("tim_business_001")
        >>>
        >>> # Upload script
        >>> add_document_to_campaign(
        ...     campaign["id"],
        ...     Path("sales_script.pdf"),
        ...     original_filename="Script Shitje Tim Business.pdf"
        ... )
        >>>
        >>> # Upload objection handling
        >>> add_document_to_campaign(
        ...     campaign["id"],
        ...     Path("objections.docx")
        ... )
        >>>
        >>> # Now campaign has 2 documents
        >>> updated = get_campaign_by_id(campaign["id"])
        >>> print(len(updated["documents"]))
        2
        >>> print(updated["total_pages"])
        15  # Combined from both PDFs

    Text Extraction:
        - PDF: Uses PyPDF2.PdfReader.extract_text()
        - DOCX: Uses python-docx Document.paragraphs
        - TXT: Direct read with UTF-8 encoding
        - Page estimation: ~300 words per page for DOCX/TXT

    Storage:
        - File: assets/campaigns/{campaign_id}/documents/{filename}
        - Text: assets/campaigns/{campaign_id}/extracted_text.txt (appended)

    Related:
        - remove_document_from_campaign() - Delete document
        - get_campaign_context_text() - Get all extracted text
        - rebuild_extracted_text() - Regenerate text file

    Note:
        - If filename conflict, adds suffix: file_1.pdf, file_2.pdf
        - Extracted text includes separator and filename header
        - Text is appended (preserves existing documents' text)
    """
    campaign = get_campaign_by_id(campaign_id)
    if not campaign:
        raise ValueError(f"Nuk u gjet fushatë me ID: {campaign_id}")

    # Kontrollo numrin e dokumenteve
    if len(campaign.get("documents", [])) >= MAX_DOCUMENTS_PER_CAMPAIGN:
        raise ValueError(f"Mund të ngarkoni maksimumi {MAX_DOCUMENTS_PER_CAMPAIGN} dokumente për fushatë")

    # Kontrollo madhësinë e file-it
    file_size_mb = file_path.stat().st_size / (1024 * 1024)
    if file_size_mb > MAX_FILE_SIZE_MB:
        raise ValueError(f"File-i është shumë i madh: {file_size_mb:.2f}MB (max: {MAX_FILE_SIZE_MB}MB)")

    # Ekstrakto tekstin
    try:
        extracted_text, num_pages = extract_text_from_file(file_path)
    except Exception as e:
        raise ValueError(f"Gabim gjatë ekstraktimit të tekstit: {e}")

    # Kontrollo totalin e faqeve
    total_pages = campaign.get("total_pages", 0) + num_pages
    if total_pages > MAX_TOTAL_PAGES:
        raise ValueError(f"Tejkalimi i limitit të faqeve: {total_pages} (max: {MAX_TOTAL_PAGES})")

    # Kopjo file-in në folderin e fushatës
    campaign_dir = CAMPAIGN_ASSETS_DIR / campaign_id / "documents"
    campaign_dir.mkdir(parents=True, exist_ok=True)

    dest_filename = original_filename or file_path.name
    dest_path = campaign_dir / dest_filename

    # Nëse ekziston file me këtë emër, shto një suffix
    counter = 1
    while dest_path.exists():
        stem = pathlib.Path(dest_filename).stem
        suffix = pathlib.Path(dest_filename).suffix
        dest_filename = f"{stem}_{counter}{suffix}"
        dest_path = campaign_dir / dest_filename
        counter += 1

    shutil.copy2(file_path, dest_path)

    # Shto metadata të dokumentit
    doc_metadata = {
        "filename": dest_filename,
        "original_name": original_filename or file_path.name,
        "uploaded_date": datetime.now().isoformat(),
        "size_kb": int(file_path.stat().st_size / 1024),
        "pages": num_pages
    }

    # Përditëso kampanjën
    campaign["documents"].append(doc_metadata)
    campaign["total_pages"] = total_pages
    campaign["total_size_kb"] = campaign.get("total_size_kb", 0) + doc_metadata["size_kb"]
    campaign["last_modified"] = datetime.now().isoformat()

    # Ruaj extracted text në një file të veçantë
    extracted_text_path = CAMPAIGN_ASSETS_DIR / campaign_id / "extracted_text.txt"

    # Append tekstin e ri
    if extracted_text_path.exists():
        with open(extracted_text_path, 'a', encoding='utf-8') as f:
            f.write(f"\n\n{'='*80}\n")
            f.write(f"DOCUMENT: {dest_filename}\n")
            f.write(f"{'='*80}\n\n")
            f.write(extracted_text)
    else:
        with open(extracted_text_path, 'w', encoding='utf-8') as f:
            f.write(f"{'='*80}\n")
            f.write(f"DOCUMENT: {dest_filename}\n")
            f.write(f"{'='*80}\n\n")
            f.write(extracted_text)

    # Ruaj në config
    config = load_campaign_config()
    for idx, c in enumerate(config["campaigns"]):
        if c["id"] == campaign_id:
            config["campaigns"][idx] = campaign
            break
    save_campaign_config(config)

    return campaign


def remove_document_from_campaign(campaign_id: str, filename: str) -> dict:
    """
    Fshin një dokument nga një fushatë

    Args:
        campaign_id: ID e fushatës
        filename: Emri i file-it që do të fshihet

    Returns:
        dict: Kampanja e përditësuar
    """
    campaign = get_campaign_by_id(campaign_id)
    if not campaign:
        raise ValueError(f"Nuk u gjet fushatë me ID: {campaign_id}")

    # Gjej dokumentin
    doc_to_remove = None
    for idx, doc in enumerate(campaign.get("documents", [])):
        if doc.get("filename") == filename:
            doc_to_remove = doc
            campaign["documents"].pop(idx)
            break

    if not doc_to_remove:
        raise ValueError(f"Nuk u gjet dokument me emrin: {filename}")

    # Fshi file-in fizik
    doc_path = CAMPAIGN_ASSETS_DIR / campaign_id / "documents" / filename
    if doc_path.exists():
        doc_path.unlink()

    # Përditëso statistikat
    campaign["total_pages"] -= doc_to_remove.get("pages", 0)
    campaign["total_size_kb"] -= doc_to_remove.get("size_kb", 0)
    campaign["last_modified"] = datetime.now().isoformat()

    # Rifillo extracted_text.txt
    rebuild_extracted_text(campaign_id)

    # Ruaj në config
    config = load_campaign_config()
    for idx, c in enumerate(config["campaigns"]):
        if c["id"] == campaign_id:
            config["campaigns"][idx] = campaign
            break
    save_campaign_config(config)

    return campaign


def rebuild_extracted_text(campaign_id: str):
    """Rikrijon file-in e tekstit të ekstraktuar nga të gjitha dokumentet"""
    campaign = get_campaign_by_id(campaign_id)
    if not campaign:
        return

    extracted_text_path = CAMPAIGN_ASSETS_DIR / campaign_id / "extracted_text.txt"

    # Fshi file-in e vjetër
    if extracted_text_path.exists():
        extracted_text_path.unlink()

    # Rikrijo nga dokumentet ekzistues
    documents_dir = CAMPAIGN_ASSETS_DIR / campaign_id / "documents"
    if not documents_dir.exists():
        return

    for doc_meta in campaign.get("documents", []):
        doc_path = documents_dir / doc_meta["filename"]
        if doc_path.exists():
            try:
                text, _ = extract_text_from_file(doc_path)
                if extracted_text_path.exists():
                    with open(extracted_text_path, 'a', encoding='utf-8') as f:
                        f.write(f"\n\n{'='*80}\n")
                        f.write(f"DOCUMENT: {doc_meta['filename']}\n")
                        f.write(f"{'='*80}\n\n")
                        f.write(text)
                else:
                    with open(extracted_text_path, 'w', encoding='utf-8') as f:
                        f.write(f"{'='*80}\n")
                        f.write(f"DOCUMENT: {doc_meta['filename']}\n")
                        f.write(f"{'='*80}\n\n")
                        f.write(text)
            except:
                pass  # Skip në rast gabimi


def get_campaign_context_text(campaign_id: str) -> str:
    """
    Merr të gjithë kontekstin e tekstit për një fushatë (dokumente të ekstraktuara)

    Returns:
        str: Teksti i kombinuar i të gjitha dokumenteve
    """
    extracted_text_path = CAMPAIGN_ASSETS_DIR / campaign_id / "extracted_text.txt"

    if extracted_text_path.exists():
        with open(extracted_text_path, 'r', encoding='utf-8') as f:
            return f.read()

    return ""


def get_campaign_hints(campaign_id: Optional[str]) -> dict:
    """
    Merr të gjithë kontekstin për një fushatë (hints + documents).

    Returns all context needed for LLM analysis: business context,
    analysis instructions, and full text from all uploaded documents.
    This is the main function called by analysis_llm.py.

    Args:
        campaign_id: ID e fushatës (e.g., "google_reserve_uk_abc123")
            If None or "", returns empty dict (no context)

    Returns:
        dict: Complete campaign context with keys:
            {
                "project_context_hint": str,  # Business/campaign context
                "summary_hint": str,          # Summary instructions for LLM
                "bullets_hint": str,          # Bullet points instructions
                "documents_text": str         # Full extracted text from all docs
            }

            If campaign not found or campaign_id is None:
            {
                "project_context_hint": "",
                "summary_hint": "",
                "bullets_hint": "",
                "documents_text": ""
            }

    Example:
        >>> # No campaign selected
        >>> hints = get_campaign_hints(None)
        >>> print(hints)
        {'project_context_hint': '', 'summary_hint': '', ...}

        >>> # With campaign
        >>> hints = get_campaign_hints("google_reserve_uk_001")
        >>> print(hints["project_context_hint"])
        "Thirrje outbound për restorante në MB..."

        >>> print(len(hints["documents_text"]))
        15000  # Full text from all documents

        >>> # Use in analysis
        >>> from core.analysis_llm import write_outputs_and_report
        >>> hints = get_campaign_hints(selected_campaign_id)
        >>> report = write_outputs_and_report(
        ...     calls,
        ...     project_context_hint=hints["project_context_hint"],
        ...     documents_text=hints["documents_text"]
        ... )

    Document Text Format:
        The documents_text field contains extracted text with separators:

        ================================================================================
        DOCUMENT: script.pdf
        ================================================================================

        [Full text of script.pdf]

        ================================================================================
        DOCUMENT: objection_handling.docx
        ================================================================================

        [Full text of objection_handling.docx]

    Related:
        - get_campaign_by_id() - Get full campaign object
        - get_campaign_context_text() - Get only documents text
        - create_campaign() - Create campaign with hints
        - update_campaign() - Modify hints

    Note:
        - Safe to call with None/invalid ID (returns empty dict)
        - Documents text is cached in extracted_text.txt for performance
        - Used by pages/1_Pipeline_Komplet.py and pages/4_Analize_Automatike.py
    """
    if not campaign_id:
        return {
            "project_context_hint": "",
            "summary_hint": "",
            "bullets_hint": "",
            "documents_text": ""
        }

    campaign = get_campaign_by_id(campaign_id)
    if not campaign:
        return {
            "project_context_hint": "",
            "summary_hint": "",
            "bullets_hint": "",
            "documents_text": ""
        }

    return {
        "project_context_hint": campaign.get("project_context_hint", ""),
        "summary_hint": campaign.get("summary_hint", ""),
        "bullets_hint": campaign.get("bullets_hint", ""),
        "documents_text": get_campaign_context_text(campaign_id)
    }

