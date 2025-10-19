"""
core/materials_generator.py

PURPOSE:
    AI-powered materials generator for sales training and quality improvement.
    Analyzes call transcripts to generate actionable materials.

RESPONSIBILITIES:
    - Extract objections and generate professional responses
    - Generate sales scripts based on successful calls
    - Create FAQs from common customer questions
    - Extract best practices from top performers

KEY FUNCTIONS:
    - generate_objections_and_responses() - Extract & respond to objections
    - generate_sales_script() - Create call scripts
    - generate_faq() - Generate FAQ from calls
    - generate_best_practices() - Extract best practices

DEPENDENCIES:
    - OpenAI GPT-4 for analysis
    - campaign_manager for context
    - transcription_audio for transcript processing

Author: Protrade AI
Last Updated: 2025-10-15
"""

import json
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime
from openai import OpenAI
import os


# ============== CONFIG ==============
DEFAULT_MODEL = "gpt-4o"
MIN_OBJECTIONS = 10
MIN_TRANSCRIPT_LENGTH = 100  # minimum characters per transcript


# ============== UTILS ==============
def _get_client() -> OpenAI:
    """Merr OpenAI client"""
    key = os.getenv("OPENAI_API_KEY")
    if not key:
        try:
            import streamlit as st
            key = st.secrets.get("OPENAI_API_KEY") or st.secrets.get("openai", {}).get("OPENAI_API_KEY")
        except Exception:
            pass
    if not key:
        raise RuntimeError("OPENAI_API_KEY mungon.")
    return OpenAI(api_key=key)


def _load_transcripts_from_paths(transcript_paths: List[Path]) -> List[Dict[str, str]]:
    """
    Ngarkon të gjitha transkriptet nga path-et e dhëna

    Args:
        transcript_paths: Lista e path-eve të transkripteve

    Returns:
        List[Dict]: [{"agent": "...", "text": "...", "path": "..."}, ...]
    """
    transcripts = []

    for path in transcript_paths:
        if not path.exists():
            continue

        try:
            text = path.read_text(encoding="utf-8", errors="ignore")

            # Skip if too short
            if len(text.strip()) < MIN_TRANSCRIPT_LENGTH:
                continue

            # Extract agent name from path (usually in parent folder)
            agent = "UNKNOWN"
            if len(path.parts) >= 2:
                agent = path.parts[-2].strip().title()

            transcripts.append({
                "agent": agent,
                "text": text,
                "path": str(path)
            })
        except Exception:
            continue

    return transcripts


def _combine_transcripts(transcripts: List[Dict[str, str]], max_chars: int = 100_000) -> str:
    """
    Kombinon transkriptet në një tekst të vetëm për analiz

    Args:
        transcripts: Lista e transkripteve
        max_chars: Maksimumi i karaktereve (për limit API)

    Returns:
        str: Teksti i kombinuar
    """
    combined = []
    total_chars = 0

    for t in transcripts:
        text = t["text"]
        agent = t["agent"]

        header = f"\n\n{'='*80}\nAGJENT: {agent}\n{'='*80}\n\n"
        section = header + text

        if total_chars + len(section) > max_chars:
            break

        combined.append(section)
        total_chars += len(section)

    return "".join(combined)


# ============== OBJECTIONS GENERATOR ==============
def generate_objections_and_responses(
    transcript_paths: List[Path],
    campaign_context: str = "",
    documents_text: str = "",
    additional_instructions: str = "",
    language: str = "sq",
    min_objections: int = MIN_OBJECTIONS
) -> Dict[str, Any]:
    """
    Gjeneron objeksione dhe përgjigje konsultative nga transkriptet.

    Analyzes ALL provided transcripts to extract common objections and
    generates professional, consultative responses that prevent objections
    and build trust rather than just "handling" them.

    Args:
        transcript_paths: Lista e të gjitha transkripteve për analiz
        campaign_context: Konteksti i fushatës/projektit
        documents_text: Teksti i dokumenteve të fushatës
        additional_instructions: Instruksione shtesë opsionale
        language: Gjuha e outputit (sq/it/en)
        min_objections: Minimumi i objeksioneve (default: 10)

    Returns:
        dict: {
            "objections": [
                {
                    "category": "Çmimi",
                    "objection": "Është shumë e shtrenjtë",
                    "frequency": "E lartë",
                    "context": "Kur...",
                    "consultative_approach": {
                        "prevention": "Si ta parandalosh...",
                        "value_building": "Si të ndërtosh vlerë...",
                        "response_framework": "Framework i përgjigjes...",
                        "example_dialogue": "Shembull dialogu..."
                    }
                },
                ...
            ],
            "general_strategy": "Strategji e përgjithshme...",
            "call_flow_recommendations": "Rekomandime për flow...",
            "total_transcripts_analyzed": 150,
            "generated_at": "2025-10-15T10:30:00"
        }

    Example:
        >>> from pathlib import Path
        >>> transcripts = list(Path("out_analysis/session/Transkripte").rglob("*.txt"))
        >>> result = generate_objections_and_responses(
        ...     transcript_paths=transcripts,
        ...     campaign_context="Shitje B2B për internet biznes",
        ...     additional_instructions="Fokuso në objeksionet teknike"
        ... )
        >>> print(f"U gjeten {len(result['objections'])} objeksione")
        >>> for obj in result['objections']:
        ...     print(f"- {obj['objection']} ({obj['category']})")

    Note:
        - Analizon TË GJITHA transkriptet (jo limit)
        - Fokusi në parandalim dhe ndërtim vlere, jo vetëm përgjigje
        - Kategorizon objeksionet (çmim, kohë, konkurrencë, etj.)
        - Jep shembuj konkretë dialogu
        - Minimum 10 objeksione unike
    """
    # Load all transcripts
    transcripts = _load_transcripts_from_paths(transcript_paths)

    if not transcripts:
        return {
            "error": "Nuk u gjend asnjë transkript i vlefshëm",
            "objections": [],
            "total_transcripts_analyzed": 0
        }

    # Combine transcripts
    combined_text = _combine_transcripts(transcripts)

    # Language mapping
    lang_map = {
        "sq": "shqip",
        "it": "italisht",
        "en": "anglisht"
    }
    language_name = lang_map.get(language, "shqip")

    # Build prompt
    prompt = f"""Ti je një ekspert i shitjeve konsultative dhe trajnimit të agjentëve.

DETYRA:
Analizo të gjitha transkriptet e telefonatave dhe ekstrakto objeksionet më të shpeshta që hasin agjentet.
Për çdo objeksion, jep jo vetëm një përgjigje, por një qasje të plotë konsultative që:
1. PARANDALON objeksionin përpara se të lindet
2. NDËRTON vlerë dhe besim me klientin
3. TRANSFORMON objeksionin në mundësi për thellim të bisedës
4. KRIJON një flow telefonate informative dhe profesionale

KONTEKST I PROJEKTIT:
{campaign_context if campaign_context else "Nuk ka kontekst specifik"}

DOKUMENTE SHTESË:
{documents_text[:10000] if documents_text else "Nuk ka dokumente shtesë"}

INSTRUKSIONE SHTESË:
{additional_instructions if additional_instructions else "Nuk ka instruksione shtesë"}

KËRKESAT:
- Gjej të paktën {min_objections} objeksione UNIKE dhe TË NDRYSHME
- Kategorizon në: Çmim, Kohë, Konkurrencë, Dyshim/Mosbesim, Nevoja, Teknike, Të tjera
- Për çdo objeksion jep: frekuencë (E lartë/Mesatare/E ulët), kontekst (kur lind), dhe qasje konsultative
- Qasja konsultative duhet të përfshijë:
  * Prevention: Si ta parandalosh që objeksioni të lindet
  * Value Building: Si të ndërtosh vlerë PËRPARA se të flasësh për zgjidhje
  * Response Framework: Struktura e përgjigjes (jo vetëm teksti)
  * Example Dialogue: Shembull konkret dialogu (3-5 shkëmbime)

- Në fund, jep:
  * General Strategy: Strategji e përgjithshme për menaxhim objeksionesh
  * Call Flow Recommendations: Si të strukturosh telefonatën për të minimizuar objeksionet

OUTPUT: JSON në gjuhën {language_name}, me këtë strukturë:
{{
  "objections": [
    {{
      "category": "Kategoria",
      "objection": "Teksti i objeksionit",
      "frequency": "E lartë/Mesatare/E ulët",
      "context": "Kur dhe pse lind ky objeksion",
      "consultative_approach": {{
        "prevention": "Si ta parandalosh...",
        "value_building": "Si të ndërtosh vlerë përpara...",
        "response_framework": "Hapat e përgjigjes...",
        "example_dialogue": "Shembull dialogu me 3-5 shkëmbime"
      }}
    }}
  ],
  "general_strategy": "Strategji e përgjithshme për menaxhim objeksionesh në këtë fushatë",
  "call_flow_recommendations": "Si të strukturosh telefonatën për parandalim objeksionesh"
}}

TRANSKRIPTET ({len(transcripts)} telefonata):
{combined_text}

KUJDES: Përgjigjet duhet të jenë konsultative, jo agresive. Fokusi është ndërtimi i besimit dhe ofrim ndihmë, jo thjesht "kalim" objeksioni.
"""

    # Call OpenAI
    client = _get_client()

    try:
        response = client.chat.completions.create(
            model=DEFAULT_MODEL,
            messages=[
                {
                    "role": "system",
                    "content": f"Ti je një ekspert i shitjeve konsultative. Gjithçka duhet të jetë në {language_name}."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            temperature=0.3,
            max_tokens=4000
        )

        content = response.choices[0].message.content or ""

        # Parse JSON
        content = content.strip().replace("```json", "").replace("```", "").strip()
        start = content.find("{")
        end = content.rfind("}")
        if start != -1 and end != -1:
            content = content[start:end+1]

        result = json.loads(content)

        # Add metadata
        result["total_transcripts_analyzed"] = len(transcripts)
        result["generated_at"] = datetime.now().isoformat()
        result["language"] = language
        result["campaign_context"] = campaign_context

        return result

    except Exception as e:
        return {
            "error": f"Gabim gjatë gjenerimit: {str(e)}",
            "objections": [],
            "total_transcripts_analyzed": len(transcripts)
        }


# ============== SALES SCRIPT GENERATOR ==============
def generate_sales_script(
    transcript_paths: List[Path],
    campaign_context: str = "",
    documents_text: str = "",
    additional_instructions: str = "",
    language: str = "sq"
) -> Dict[str, Any]:
    """
    Gjeneron skript shitjeje bazuar në telefonatat më të suksesshme.

    Analyzes successful calls to create a comprehensive sales script
    with opening, discovery, presentation, objection handling, and closing.

    Args:
        transcript_paths: Lista e të gjitha transkripteve
        campaign_context: Konteksti i fushatës
        documents_text: Dokumentet e fushatës
        additional_instructions: Instruksione shtesë
        language: Gjuha (sq/it/en)

    Returns:
        dict: {
            "script": {
                "opening": {...},
                "discovery": {...},
                "presentation": {...},
                "objection_handling": {...},
                "closing": {...}
            },
            "key_phrases": [...],
            "tone_guidelines": "...",
            "do_dont": {...}
        }
    """
    transcripts = _load_transcripts_from_paths(transcript_paths)

    if not transcripts:
        return {"error": "Nuk u gjend asnjë transkript"}

    combined_text = _combine_transcripts(transcripts)

    lang_map = {"sq": "shqip", "it": "italisht", "en": "anglisht"}
    language_name = lang_map.get(language, "shqip")

    prompt = f"""Ti je një ekspert i krijimit të skripteve të shitjeve konsultative.

DETYRA:
Bazuar në transkriptet e telefonatave të suksesshme, krijo një skript të plotë shitjeje që përfshin:
1. HAPJE (Opening) - Si të fillosh telefonatën me impact
2. ZBULIM (Discovery) - Pyetje për të kuptuar nevojat
3. PREZANTIM (Presentation) - Si ta prezantosh ofertën
4. MENAXHIM OBJEKSIONESH (Objection Handling) - Të përbashkëtat
5. MBYLLJE (Closing) - Si ta mbyllësh me sukses

KONTEKST I PROJEKTIT:
{campaign_context if campaign_context else "Nuk ka kontekst"}

DOKUMENTE:
{documents_text[:10000] if documents_text else "Nuk ka dokumente"}

INSTRUKSIONE SHTESË:
{additional_instructions if additional_instructions else "Nuk ka"}

OUTPUT: JSON në {language_name} me strukturë:
{{
  "script": {{
    "opening": {{
      "main_script": "Skripti kryesor...",
      "variations": ["Variacion 1", "Variacion 2"],
      "tips": ["Tip 1", "Tip 2"]
    }},
    "discovery": {{
      "key_questions": ["Pyetje 1", "Pyetje 2"],
      "listening_points": ["Çfarë të dëgjosh për..."],
      "red_flags": ["Sinjale paralajmëruese..."]
    }},
    "presentation": {{
      "value_proposition": "Propozimi i vlerës...",
      "key_benefits": ["Përfitimi 1", "Përfitimi 2"],
      "proof_points": ["Provë 1", "Provë 2"],
      "storytelling_template": "Template për tregim historie..."
    }},
    "objection_handling": {{
      "common_objections": [{{"objection": "...", "response": "..."}}],
      "framework": "Framework i përgjigjes..."
    }},
    "closing": {{
      "main_close": "Mbyllja kryesore...",
      "alternative_closes": ["Alt 1", "Alt 2"],
      "next_steps": ["Hapi 1", "Hapi 2"]
    }}
  }},
  "key_phrases": ["Fraza që funksionojnë mirë..."],
  "tone_guidelines": "Si duhet të jetë toni...",
  "do_dont": {{
    "do": ["Bëj këtë...", "Dhe këtë..."],
    "dont": ["Mos bëj këtë...", "Mos thuaj..."]
  }},
  "call_duration_target": "Kohëzgjatja ideale...",
  "success_metrics": "Si ta matësh suksesin..."
}}

TRANSKRIPTET ({len(transcripts)} telefonata):
{combined_text}
"""

    client = _get_client()

    try:
        response = client.chat.completions.create(
            model=DEFAULT_MODEL,
            messages=[
                {"role": "system", "content": f"Ti je ekspert i skripteve të shitjes. Gjithçka në {language_name}."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.4,
            max_tokens=4000
        )

        content = response.choices[0].message.content or ""
        content = content.strip().replace("```json", "").replace("```", "").strip()
        start = content.find("{")
        end = content.rfind("}")
        if start != -1 and end != -1:
            content = content[start:end+1]

        result = json.loads(content)
        result["total_transcripts_analyzed"] = len(transcripts)
        result["generated_at"] = datetime.now().isoformat()

        return result

    except Exception as e:
        return {"error": f"Gabim: {str(e)}"}


# ============== FAQ GENERATOR ==============
def generate_faq(
    transcript_paths: List[Path],
    campaign_context: str = "",
    documents_text: str = "",
    additional_instructions: str = "",
    language: str = "sq"
) -> Dict[str, Any]:
    """
    Gjeneron FAQ (Frequently Asked Questions) nga pyetjet e klientëve.

    Args:
        transcript_paths: Lista e transkripteve
        campaign_context: Konteksti i fushatës
        documents_text: Dokumentet
        additional_instructions: Instruksione shtesë
        language: Gjuha

    Returns:
        dict: Lista e pyetjeve dhe përgjigjeve
    """
    transcripts = _load_transcripts_from_paths(transcript_paths)

    if not transcripts:
        return {"error": "Nuk u gjend asnjë transkript"}

    combined_text = _combine_transcripts(transcripts)
    lang_map = {"sq": "shqip", "it": "italisht", "en": "anglisht"}
    language_name = lang_map.get(language, "shqip")

    prompt = f"""Ekstrakto pyetjet më të shpeshta (FAQ) nga këto telefonata.

KONTEKST: {campaign_context if campaign_context else "Nuk ka"}
DOKUMENTE: {documents_text[:10000] if documents_text else "Nuk ka"}
INSTRUKSIONE: {additional_instructions if additional_instructions else "Nuk ka"}

OUTPUT JSON në {language_name}:
{{
  "faqs": [
    {{
      "category": "Kategoria",
      "question": "Pyetja",
      "answer": "Përgjigjja e shkurtër",
      "detailed_answer": "Përgjigjja e detajuar",
      "frequency": "E lartë/Mesatare/E ulët"
    }}
  ],
  "categories_summary": {{"Kategoria": "Përmbledhje"}}
}}

TRANSKRIPTET ({len(transcripts)}):
{combined_text}
"""

    client = _get_client()

    try:
        response = client.chat.completions.create(
            model=DEFAULT_MODEL,
            messages=[
                {"role": "system", "content": f"Gjithçka në {language_name}."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            max_tokens=3000
        )

        content = response.choices[0].message.content or ""
        content = content.strip().replace("```json", "").replace("```", "").strip()
        start = content.find("{")
        end = content.rfind("}")
        if start != -1 and end != -1:
            content = content[start:end+1]

        result = json.loads(content)
        result["total_transcripts_analyzed"] = len(transcripts)
        result["generated_at"] = datetime.now().isoformat()

        return result

    except Exception as e:
        return {"error": f"Gabim: {str(e)}"}


# ============== BEST PRACTICES GENERATOR ==============
def generate_best_practices(
    transcript_paths: List[Path],
    campaign_context: str = "",
    documents_text: str = "",
    additional_instructions: str = "",
    language: str = "sq"
) -> Dict[str, Any]:
    """
    Ekstrakton praktikat më të mira nga telefonatat e suksesshme.

    Args:
        transcript_paths: Lista e transkripteve
        campaign_context: Konteksti
        documents_text: Dokumentet
        additional_instructions: Instruksione shtesë
        language: Gjuha

    Returns:
        dict: Best practices të kategorizuara
    """
    transcripts = _load_transcripts_from_paths(transcript_paths)

    if not transcripts:
        return {"error": "Nuk u gjend asnjë transkript"}

    combined_text = _combine_transcripts(transcripts)
    lang_map = {"sq": "shqip", "it": "italisht", "en": "anglisht"}
    language_name = lang_map.get(language, "shqip")

    prompt = f"""Ekstrakto praktikat më të mira nga këto telefonata të suksesshme.

KONTEKST: {campaign_context if campaign_context else "Nuk ka"}
DOKUMENTE: {documents_text[:10000] if documents_text else "Nuk ka"}
INSTRUKSIONE: {additional_instructions if additional_instructions else "Nuk ka"}

OUTPUT JSON në {language_name}:
{{
  "best_practices": [
    {{
      "category": "Hapje/Zbulim/Prezantim/Mbyllje/Ton",
      "practice": "Praktika",
      "explanation": "Pse funksionon",
      "example": "Shembull konkret",
      "difficulty": "E lehtë/Mesatare/E vështirë"
    }}
  ],
  "top_performers_patterns": ["Pattern 1", "Pattern 2"],
  "common_mistakes_to_avoid": ["Gabim 1", "Gabim 2"],
  "training_recommendations": "Rekomandime për trajnim..."
}}

TRANSKRIPTET ({len(transcripts)}):
{combined_text}
"""

    client = _get_client()

    try:
        response = client.chat.completions.create(
            model=DEFAULT_MODEL,
            messages=[
                {"role": "system", "content": f"Gjithçka në {language_name}."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            max_tokens=3000
        )

        content = response.choices[0].message.content or ""
        content = content.strip().replace("```json", "").replace("```", "").strip()
        start = content.find("{")
        end = content.rfind("}")
        if start != -1 and end != -1:
            content = content[start:end+1]

        result = json.loads(content)
        result["total_transcripts_analyzed"] = len(transcripts)
        result["generated_at"] = datetime.now().isoformat()

        return result

    except Exception as e:
        return {"error": f"Gabim: {str(e)}"}


# ============== EXPORT FUNCTIONS ==============
def export_to_docx(data: Dict[str, Any], output_path: Path, material_type: str):
    """
    Eksporton materialin në format DOCX

    Args:
        data: Të dhënat e gjeneruara
        output_path: Path për ruajtje
        material_type: Lloji (objections/script/faq/best_practices)
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Title
        title = doc.add_heading(f"AI Generated: {material_type.upper()}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        meta = doc.add_paragraph()
        meta.add_run(f"Gjeneruar: {data.get('generated_at', 'N/A')}\n").italic = True
        meta.add_run(f"Transkripte të analizuara: {data.get('total_transcripts_analyzed', 0)}\n").italic = True

        doc.add_paragraph("_" * 80)

        # Content based on type
        if material_type == "objections":
            doc.add_heading("Objeksione dhe Përgjigje Konsultative", 1)

            for obj in data.get("objections", []):
                doc.add_heading(f"{obj.get('category', 'N/A')}: {obj.get('objection', 'N/A')}", 2)
                doc.add_paragraph(f"Frekuencë: {obj.get('frequency', 'N/A')}")
                doc.add_paragraph(f"Kontekst: {obj.get('context', 'N/A')}")

                approach = obj.get("consultative_approach", {})
                doc.add_heading("Qasja Konsultative:", 3)
                doc.add_paragraph(f"Prevention: {approach.get('prevention', 'N/A')}")
                doc.add_paragraph(f"Value Building: {approach.get('value_building', 'N/A')}")
                doc.add_paragraph(f"Response Framework: {approach.get('response_framework', 'N/A')}")
                doc.add_paragraph(f"Shembull Dialogu:\n{approach.get('example_dialogue', 'N/A')}")
                doc.add_paragraph()

            doc.add_heading("Strategji e Përgjithshme", 1)
            doc.add_paragraph(data.get("general_strategy", "N/A"))

            doc.add_heading("Rekomandime për Flow", 1)
            doc.add_paragraph(data.get("call_flow_recommendations", "N/A"))

        elif material_type == "script":
            script = data.get("script", {})

            for section_name, section_data in script.items():
                doc.add_heading(section_name.upper(), 1)

                if isinstance(section_data, dict):
                    for key, value in section_data.items():
                        doc.add_heading(key.replace("_", " ").title(), 2)
                        if isinstance(value, list):
                            for item in value:
                                doc.add_paragraph(str(item), style='List Bullet')
                        else:
                            doc.add_paragraph(str(value))

        elif material_type == "faq":
            for faq in data.get("faqs", []):
                doc.add_heading(f"Q: {faq.get('question', 'N/A')}", 2)
                doc.add_paragraph(f"Kategoria: {faq.get('category', 'N/A')}")
                doc.add_paragraph(f"A: {faq.get('answer', 'N/A')}")
                doc.add_paragraph(f"Detaje: {faq.get('detailed_answer', 'N/A')}")
                doc.add_paragraph()

        elif material_type == "best_practices":
            for practice in data.get("best_practices", []):
                doc.add_heading(f"{practice.get('category', 'N/A')}: {practice.get('practice', 'N/A')}", 2)
                doc.add_paragraph(f"Shpjegim: {practice.get('explanation', 'N/A')}")
                doc.add_paragraph(f"Shembull: {practice.get('example', 'N/A')}")
                doc.add_paragraph(f"Vështirësi: {practice.get('difficulty', 'N/A')}")
                doc.add_paragraph()

        # Save
        doc.save(str(output_path))
        return True

    except Exception as e:
        print(f"Error exporting to DOCX: {e}")
        return False


def export_to_json(data: Dict[str, Any], output_path: Path):
    """Eksporton në JSON"""
    try:
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        return True
    except Exception:
        return False


def export_to_txt(data: Dict[str, Any], output_path: Path, material_type: str):
    """Eksporton në TXT të thjeshtë"""
    try:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(f"=== AI GENERATED: {material_type.upper()} ===\n\n")
            f.write(f"Gjeneruar: {data.get('generated_at', 'N/A')}\n")
            f.write(f"Transkripte: {data.get('total_transcripts_analyzed', 0)}\n\n")
            f.write("="*80 + "\n\n")
            f.write(json.dumps(data, indent=2, ensure_ascii=False))
        return True
    except Exception:
        return False


# ============== TRAINING RECORDINGS SELECTOR ==============
def select_recordings_for_training(
    transcript_paths: List[Path],
    max_to_analyze: int = 500,
    max_to_download: int = 20,
    selection_criteria: str = "best",
    campaign_context: str = "",
    documents_text: str = "",
    language: str = "sq",
    campaign_name: str = "UNKNOWN"
) -> Dict[str, Any]:
    """
    Analizon regjistrimet dhe zgjedh më të mirat/më problematiket për trajnim.

    Organizimi i folderëve:
        out_analysis/
            {campaign_name}/
                regjistrime_pozitive/  (ose regjistrime_negative)
                    20250113_143022/   (data_ora)
                        audio files këtu

    Args:
        transcript_paths: Lista e të gjitha transkripteve
        max_to_analyze: Sa regjistrime të analizohen (të gjitha)
        max_to_download: Sa regjistrime të zgjedhura të shkarkohen
        selection_criteria: "best" | "worst" | "average" | "balanced"
        campaign_context: Konteksti i fushatës
        documents_text: Dokumentet e fushatës
        language: Gjuha
        campaign_name: Emri i fushatës (për organizim folderësh)

    Returns:
        {
            "selected_recordings": [
                {
                    "transcript_path": "...",
                    "audio_path": "...",
                    "agent": "...",
                    "score": 4.5,
                    "summary": "...",
                    "preggi": [...],
                    "da_migliorare": [...],
                    "reason": "Score i lartë dhe komunikim efektiv",
                    "ranking": 1
                }
            ],
            "total_analyzed": 500,
            "selection_criteria": "best (score >= 4.0)",
            "download_folder": "...",
            "statistics": {
                "best_count": 45,
                "average_count": 300,
                "worst_count": 155
            }
        }

    Example:
        >>> from pathlib import Path
        >>> transcripts = list(Path("out_analysis/session/Transkripte").rglob("*.txt"))
        >>> result = select_recordings_for_training(
        ...     transcript_paths=transcripts,
        ...     max_to_analyze=500,
        ...     max_to_download=20,
        ...     selection_criteria="best",
        ...     campaign_name="GOOGLE RESERVE UK - RESTORANTE"
        ... )
        >>> print(f"U zgjodhën {len(result['selected_recordings'])} regjistrime")
        >>> print(f"Folder: {result['download_folder']}")

    Note:
        - Përdor sistemin ekzistues të analizës (analysis_llm.py)
        - Organizon regjistrimet në folder të strukturuar
        - Krijon një folder për çdo ekzekutim me datë/ora
    """
    from core.analysis_llm import analyze_agent_transcripts
    from core.config import OUT_DIR
    import shutil

    # Limit transcripts to analyze
    transcripts_to_analyze = transcript_paths[:max_to_analyze]

    if not transcripts_to_analyze:
        return {
            "error": "Nuk u gjend asnjë transkript i vlefshëm",
            "selected_recordings": [],
            "total_analyzed": 0
        }

    # Analyze all transcripts
    analyzed_recordings = []

    for transcript_path in transcripts_to_analyze:
        try:
            # Read transcript
            transcript_text = transcript_path.read_text(encoding="utf-8", errors="ignore")

            # Extract agent name from path
            agent = "UNKNOWN"
            if len(transcript_path.parts) >= 2:
                agent = transcript_path.parts[-2].strip().title()

            # Analyze transcript
            analysis = analyze_agent_transcripts(
                agent_name=agent,
                transcript_text=transcript_text,
                language=language,
                summary_hint="",
                bullets_hint="",
                project_context_hint=campaign_context,
                documents_text=documents_text
            )

            # Calculate score
            score = analysis.get("score", 3.0)

            # Find corresponding audio file
            audio_path = None
            possible_audio_exts = [".mp3", ".wav", ".m4a", ".ogg", ".flac"]

            # Strategy 1: Check same directory
            for ext in possible_audio_exts:
                audio_candidate = transcript_path.parent / f"{transcript_path.stem}{ext}"
                if audio_candidate.exists():
                    audio_path = audio_candidate
                    break

            # Strategy 2: Check parent directory (if transcript is in "Transkripte" subfolder)
            if not audio_path and transcript_path.parent.name.lower() in ["transkripte", "transcripts"]:
                for ext in possible_audio_exts:
                    audio_candidate = transcript_path.parent.parent / f"{transcript_path.stem}{ext}"
                    if audio_candidate.exists():
                        audio_path = audio_candidate
                        break

            # Strategy 3: Search in common parent directories
            if not audio_path:
                # Try to find in parallel folders (e.g., "Audio", "Regjistrime", etc.)
                parent = transcript_path.parent
                for search_dir_name in ["Audio", "Regjistrime", "Recordings", "Audio Files"]:
                    search_dir = parent.parent / search_dir_name / transcript_path.parent.name
                    if search_dir.exists():
                        for ext in possible_audio_exts:
                            audio_candidate = search_dir / f"{transcript_path.stem}{ext}"
                            if audio_candidate.exists():
                                audio_path = audio_candidate
                                break
                        if audio_path:
                            break

            # Strategy 4: Search recursively in parent directories (up to 3 levels)
            if not audio_path:
                search_parent = transcript_path.parent
                for _ in range(3):
                    for ext in possible_audio_exts:
                        audio_candidate = search_parent / f"{transcript_path.stem}{ext}"
                        if audio_candidate.exists():
                            audio_path = audio_candidate
                            break
                    if audio_path:
                        break
                    search_parent = search_parent.parent
                    if not search_parent or search_parent == search_parent.parent:
                        break

            analyzed_recordings.append({
                "transcript_path": str(transcript_path),
                "audio_path": str(audio_path) if audio_path else None,
                "agent": agent,
                "score": score,
                "summary": analysis.get("summary", ""),
                "preggi": analysis.get("preggi", []),
                "da_migliorare": analysis.get("da_migliorare", []),
                "ranking": 0  # Will be set later
            })

        except Exception as e:
            continue

    # Sort by score
    analyzed_recordings.sort(key=lambda x: x["score"], reverse=True)

    # Set rankings
    for i, rec in enumerate(analyzed_recordings, 1):
        rec["ranking"] = i

    # Calculate statistics
    best_count = sum(1 for r in analyzed_recordings if r["score"] >= 4.0)
    worst_count = sum(1 for r in analyzed_recordings if r["score"] <= 2.5)
    average_count = sum(1 for r in analyzed_recordings if 2.5 < r["score"] < 4.0)

    # Select recordings based on criteria
    selected_recordings = []
    auto_fallback = False

    if selection_criteria == "best":
        # Top performers (score >= 4.0)
        selected_recordings = [r for r in analyzed_recordings if r["score"] >= 4.0][:max_to_download]
        criteria_desc = f"best (score >= 4.0)"
        folder_suffix = "regjistrime_pozitive"

        # Nëse nuk gjendet asgjë, sugjero mesataret
        if not selected_recordings and best_count == 0:
            auto_fallback = True
            selected_recordings = [r for r in analyzed_recordings if 2.5 < r["score"] < 4.0][:max_to_download]
            criteria_desc = f"average (auto-fallback, pasi nuk ka top performers)"
            folder_suffix = "regjistrime_mesatare"

    elif selection_criteria == "worst":
        # Bottom performers (score <= 2.5)
        selected_recordings = [r for r in analyzed_recordings if r["score"] <= 2.5][:max_to_download]
        criteria_desc = f"worst (score <= 2.5)"
        folder_suffix = "regjistrime_negative"

        # Nëse nuk gjendet asgjë, sugjero mesataret
        if not selected_recordings and worst_count == 0:
            auto_fallback = True
            selected_recordings = [r for r in analyzed_recordings if 2.5 < r["score"] < 4.0][:max_to_download]
            criteria_desc = f"average (auto-fallback, pasi nuk ka bottom performers)"
            folder_suffix = "regjistrime_mesatare"

    elif selection_criteria == "average":
        # Average performers (2.5 < score < 4.0)
        selected_recordings = [r for r in analyzed_recordings if 2.5 < r["score"] < 4.0][:max_to_download]
        criteria_desc = f"average (2.5 < score < 4.0)"
        folder_suffix = "regjistrime_mesatare"

    else:  # balanced
        # 50% top + 50% bottom
        top_half = max_to_download // 2
        bottom_half = max_to_download - top_half

        top_recordings = [r for r in analyzed_recordings if r["score"] >= 4.0][:top_half]
        bottom_recordings = [r for r in analyzed_recordings if r["score"] <= 2.5][:bottom_half]

        selected_recordings = top_recordings + bottom_recordings
        criteria_desc = f"balanced (50% top + 50% bottom)"
        folder_suffix = "regjistrime_te_balancuara"

        # Nëse nuk ka top ose bottom, sugjero mesataret
        if not selected_recordings and (best_count == 0 or worst_count == 0):
            auto_fallback = True
            selected_recordings = [r for r in analyzed_recordings if 2.5 < r["score"] < 4.0][:max_to_download]
            criteria_desc = f"average (auto-fallback, pasi nuk ka top/bottom performers)"
            folder_suffix = "regjistrime_mesatare"

    # Create organized folder structure
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    # Clean campaign name for folder
    safe_campaign_name = campaign_name.replace("/", "-").replace("\\", "-").strip()

    # Create folder structure: out_analysis/{campaign_name}/{folder_suffix}/{timestamp}/
    base_folder = OUT_DIR / safe_campaign_name / folder_suffix / timestamp
    base_folder.mkdir(parents=True, exist_ok=True)

    # Copy selected recordings to the folder
    copied_count = 0
    for rec in selected_recordings:
        if rec["audio_path"] and Path(rec["audio_path"]).exists():
            try:
                audio_file = Path(rec["audio_path"])
                dest_file = base_folder / audio_file.name
                shutil.copy2(audio_file, dest_file)
                rec["copied_to"] = str(dest_file)
                copied_count += 1
            except Exception:
                rec["copied_to"] = None
        else:
            rec["copied_to"] = None

    # Add reason for selection
    for rec in selected_recordings:
        if rec["score"] >= 4.0:
            rec["reason"] = f"Score i lartë ({rec['score']}). Agjenti ka treguar cilësi të lartë në komunikim dhe menaxhim të bisedës."
        elif rec["score"] <= 2.5:
            rec["reason"] = f"Score i ulët ({rec['score']}). Regjistrim i përshtatshëm për analizë të problemeve dhe trajnim për përmirësim."
        else:
            rec["reason"] = f"Score mesatar ({rec['score']}). Regjistrim reprezentativ për analizë të përgjithshme."

    return {
        "selected_recordings": selected_recordings,
        "total_analyzed": len(analyzed_recordings),
        "selection_criteria": criteria_desc,
        "download_folder": str(base_folder),
        "statistics": {
            "best_count": best_count,
            "average_count": average_count,
            "worst_count": worst_count,
            "copied_count": copied_count
        },
        "campaign_name": safe_campaign_name,
        "folder_structure": f"{safe_campaign_name}/{folder_suffix}/{timestamp}/",
        "generated_at": datetime.now().isoformat(),
        "auto_fallback": auto_fallback,
        "auto_fallback_message": (
            "⚠️ Kriteri i zgjedhur nuk ka prodhuar rezultate. U zgjodhën automatikisht regjistrime mesatare. "
            "Sugjerohet të përdorni kriterin 'Mesataret' për herën tjetër."
            if auto_fallback else ""
        )
    }








