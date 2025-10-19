"""
pages/2_Tools.py
Mjete të Ndara - Filtrim, Upload, Transkriptim
Version: v2.0
Author: Protrade AI
"""

import streamlit as st

st.title("🛠️ Mjete të Ndara")
st.caption("Përdor këto mjete kur ke nevojë të ekzekutosh hapa individualë jashtë pipeline-it kryesor.")

# Create tabs
tab1, tab2, tab3, tab4 = st.tabs(["📥 Filtrim & Shkarkim", "☁️ Drive Upload", "📝 Transkriptim", "🤖 Materiale AI"])

# ======================== TAB 1: FILTRIM & SHKARKIM ========================
with tab1:
    import pathlib
    from datetime import datetime, time
    from core.db_vicidial import list_recordings, set_db_connection, get_current_db_key, _read_db_secrets
    from core.downloader_vicidial import download_recording
    from core.config import OUT_DIR
    import urllib3
    urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

    st.markdown("### 📥 Filtrim & Shkarkim (Vicidial DB → Lokale)")
    st.caption("Merr regjistrimet nga `recording_log` dhe i ruan lokalisht.")

    # Database selector
    col_db1, col_db2 = st.columns([3, 1])
    with col_db1:
        db1_host, _, _, _ = _read_db_secrets("db")
        db2_host, _, _, _ = _read_db_secrets("db2")
        db_options = {
            f"Database 1 ({db1_host or 'Default'})": "db",
            f"Database 2 ({db2_host or '95.217.87.125'})": "db2"
        }
        db_label = st.selectbox(
            "Zgjidh Databazën",
            options=list(db_options.keys()),
            index=0 if get_current_db_key() == "db" else 1,
            help="Zgjidh nga cila databazë të lexohen të dhënat",
            key="tools_db_select"
        )
        selected_db_key = db_options[db_label]
        set_db_connection(selected_db_key)
    with col_db2:
        current_host, _, _, _ = _read_db_secrets(selected_db_key)
        st.caption(f"🔌 **{current_host}**")

    st.markdown("---")

    # Filtrat
    col_date1, col_date2 = st.columns(2)
    with col_date1:
        start_date = st.date_input("Data fillimit", key="tools_start_date")
        start_time = st.time_input("Ora fillimit", value=time(0,0,0), key="tools_start_time")
    with col_date2:
        end_date = st.date_input("Data mbarimit", key="tools_end_date")
        end_time = st.time_input("Ora mbarimit", value=time(23,59,59), key="tools_end_time")

    session_name = st.text_input("Emri i Session-it", key="tools_session")
    campaign = st.text_input("Fushata (Campaign) - opsional", key="tools_campaign", help="Filtron join-in me vicidial_log")
    max_files = st.number_input("Maksimumi i skedarëve për shkarkim", 1, 100000, 500, key="tools_max_files")

    # Filtrimi i kohëzgjatjes
    st.markdown("**Filtrimi i kohëzgjatjes (opsional)**")
    col_dur1, col_dur2 = st.columns(2)
    with col_dur1:
        min_duration = st.number_input("Kohëzgjatja minimale (sekonda)", 0, 3600, 0, help="0 = pa kufi", key="tools_min_dur")
    with col_dur2:
        max_duration = st.number_input("Kohëzgjatja maksimale (sekonda)", 0, 3600, 0, help="0 = pa kufi", key="tools_max_dur")

    # Basic auth
    col_auth1, col_auth2 = st.columns(2)
    with col_auth1:
        basic_user = st.text_input("Basic auth user (nëse duhet)", key="tools_auth_user")
    with col_auth2:
        basic_pass = st.text_input("Basic auth pass (nëse duhet)", type="password", key="tools_auth_pass")

    run_download = st.button("⬇️ Shkarko regjistrimet", type="primary", disabled=not session_name, key="tools_download_btn")

    if run_download:
        start_dt = datetime.combine(start_date, start_time).strftime("%Y-%m-%d %H:%M:%S")
        end_dt = datetime.combine(end_date, end_time).strftime("%Y-%m-%d %H:%M:%S")
        st.info(f"Po kërkoj regjistrimet nga DB: {start_dt} → {end_dt} | campaign={campaign or '(të gjitha)'}")

        try:
            rows = list_recordings(start_dt, end_dt, campaign.strip() or None, limit=int(max_files))
        except Exception as e:
            st.error(f"Gabim gjatë query: {e}")
            st.stop()

        if not rows:
            st.warning("S'u gjet asnjë regjistrim në këtë interval/filtra.")
            st.stop()

        OUT_DIR.mkdir(parents=True, exist_ok=True)
        root = OUT_DIR / session_name
        downloaded = 0
        failed = 0
        filtered_count = 0

        prog = st.progress(0, text="Duke shkarkuar...")
        total = len(rows)

        for i, r in enumerate(rows, start=1):
            user = (r.get("user") or "UNKNOWN").strip().title()
            campaign_id = (r.get("campaign_id") or (campaign or "UNKNOWN")).strip()
            filename = r.get("filename") or "recording"
            location = r.get("location") or ""
            length_sec = r.get("length_in_sec") or 0

            # Filtrim kohëzgjatje
            if min_duration > 0 and length_sec < min_duration:
                filtered_count += 1
                continue
            if max_duration > 0 and length_sec > max_duration:
                filtered_count += 1
                continue

            ext = pathlib.Path(location).suffix or ".wav"
            safe_user = user.replace("/", "-")
            safe_campaign = campaign_id.replace("/", "-")

            out_dir = root / safe_campaign / safe_user
            out_file = out_dir / f"{filename}{ext}"

            try:
                ok = download_recording(location, out_file, auth=(basic_user, basic_pass) if basic_user or basic_pass else None)
                if ok:
                    downloaded += 1
                else:
                    failed += 1
            except Exception as e:
                failed += 1
                st.write(f"🚫 {filename}: {e}")

            prog.progress(int(i/total*100), text=f"Shkarkim: {i}/{total}")

            if downloaded >= max_files:
                break

        if filtered_count > 0:
            st.info(f"ℹ️ U filtruan {filtered_count} regjistrime për shkak të kohëzgjatjes.")

        st.success(f"✅ Sukses: {downloaded} • ❌ Dështime: {failed}")
        st.info(f"📁 Skedarët janë ruajtur te: `{root}`")

# ======================== TAB 2: DRIVE UPLOAD ========================
with tab2:
    import mimetypes
    import pathlib
    from collections import defaultdict
    from typing import Dict, List
    from googleapiclient.discovery import build
    from core.drive_io import get_user_oauth_creds, ensure_path, upload_file, force_reauth

    st.markdown("### ☁️ Drive Upload")
    st.caption("Struktura: <Parent>/<Session>/<Kampanja>/<Agjent>/…")

    # Parent ID
    PARENT_ID = None
    try:
        PARENT_ID = st.secrets.get("drive", {}).get("parent_id")
    except Exception:
        pass
    if not PARENT_ID:
        try:
            from core.config import DRIVE_PARENT_ID as _PID
            PARENT_ID = _PID
        except Exception:
            PARENT_ID = ""

    col_upload1, col_upload2 = st.columns(2)
    with col_upload1:
        local_root = st.text_input("Folder lokal me regjistrime (audio)", value=str(pathlib.Path.cwd()), key="upload_local_root")
    with col_upload2:
        parent_id = st.text_input("Parent Folder ID në Drive", value=PARENT_ID or "", help="ID e folderit ku do krijohet Session/Campaign/Agent", key="upload_parent_id")

    session_name_upload = st.text_input("Emri i Folderit kryesor (Session)", placeholder="p.sh. 2025-10-06_Batch1", key="upload_session")
    campaign_name_upload = st.text_input("Emri i fushatës (Campaign)", placeholder="p.sh. AUTOBIZ", key="upload_campaign")

    col_btn1, col_btn2 = st.columns(2)
    with col_btn1:
        run_upload = st.button("📤 Ngarko në Drive", type="primary", disabled=not (local_root and parent_id and session_name_upload and campaign_name_upload), key="upload_btn")
    with col_btn2:
        if st.button("🔄 Re-link Google Drive", key="upload_reauth_btn"):
            try:
                force_reauth(readonly=False)
                st.success("✅ Relidhja u krye.")
            except Exception as e:
                st.error(f"❌ Re-auth dështoi: {e}")

    BAD = {"new folder","downloads","desktop","vicidial_agent","data","out_analysis","audio","records","regjistrime"}

    def guess_agent_from_path(p: pathlib.Path) -> str:
        for seg in reversed(p.parts):
            s = seg.replace("_"," ").strip()
            if "." in s:
                continue
            words = s.split()
            if 1 <= len(words) <= 3 and s.lower() not in BAD:
                return s.title()
        return "UNKNOWN"

    if run_upload:
        root = pathlib.Path(local_root)
        if not root.exists():
            st.error("❌ Folderi lokal nuk ekziston.")
            st.stop()

        audio_exts = {".mp3",".wav",".m4a",".mp4",".ogg",".flac",".txt",".docx"}
        files = [p for p in root.rglob("*") if p.suffix.lower() in audio_exts]

        if not files:
            st.warning("⚠️ S'u gjet asnjë file audio/transkriptim në atë folder.")
            st.stop()

        creds = get_user_oauth_creds(readonly=False)
        drive = build("drive", "v3", credentials=creds)

        session_id = ensure_path(drive, parent_id, [session_name_upload])
        campaign_id = ensure_path(drive, session_id, [campaign_name_upload])

        grouped: Dict[str, List[pathlib.Path]] = defaultdict(list)
        for p in files:
            agent = guess_agent_from_path(p)
            grouped[agent].append(p)

        total = sum(len(v) for v in grouped.values())
        prog = st.progress(0, text="Duke ngarkuar...")
        done = 0

        for agent, plist in grouped.items():
            agent_folder_id = ensure_path(drive, campaign_id, [agent])
            for p in plist:
                mime, _ = mimetypes.guess_type(str(p))
                try:
                    upload_file(drive, agent_folder_id, str(p), mime_type=mime)
                    done += 1
                    percent = int(done/total*100)
                    prog.progress(percent, text=f"Ngarkim: {done}/{total}")
                except Exception as e:
                    st.write(f"🚫 {p.name}: {e}")

        st.success(f"✅ U ngarkuan {done}/{total} file në Drive te {session_name_upload}/{campaign_name_upload}/<Agent>/")

# ======================== TAB 3: TRANSKRIPTIM ========================
with tab3:
    import pathlib, json
    from core.transcription_audio import transcribe_audio_files
    from core.config import OUT_DIR

    st.markdown("### 📝 Transkriptim (Audio → TXT/DOCX)")
    st.caption("Strategjia: Direct-first (4o-transcribe) → fallback WAV → fallback whisper-1")

    # Parametrat
    session_name_transcribe = st.text_input("Emri i sesionit (nën out_analysis/)", placeholder="p.sh. 2025-10-07_Batch1", key="transcribe_session")
    mode = st.radio("Mode", ["per-file", "merged"], index=0, horizontal=True, key="transcribe_mode")

    col_trans1, col_trans2, col_trans3 = st.columns(3)
    with col_trans1:
        reuse_existing = st.checkbox("Përdor caching", value=True, key="transcribe_cache")
    with col_trans2:
        save_docx = st.checkbox("Ruaj edhe .docx", value=False, key="transcribe_docx")
    with col_trans3:
        force = st.checkbox("Ritranskripto nga e para (force)", value=False, key="transcribe_force")

    # Organizimi sipas agjentit
    st.markdown("**Organizimi sipas agjentit (opsional)**")
    organize_by_agent = st.checkbox("Organizo transkriptet sipas emrit të agjentit", value=False, key="transcribe_organize",
                                   help="Krijo folder të ndara për çdo agjent në Transkripte/")

    agent_mapping = None
    if organize_by_agent:
        st.caption("Përcakto emrat e agjentëve për file-a (opsional)")
        agent_mapping_text = st.text_area(
            "Mapping agjentësh (format: filename1=Agent1, filename2=Agent2)",
            placeholder="call_001.mp3=John Doe, call_002.mp3=Jane Smith",
            height=100,
            key="transcribe_mapping"
        )

        if agent_mapping_text.strip():
            agent_mapping = {}
            for mapping in agent_mapping_text.split(','):
                if '=' in mapping:
                    filename, agent = mapping.strip().split('=', 1)
                    agent_mapping[filename.strip()] = agent.strip()

    # Upload audio
    st.markdown("**Ngarko audio**")
    upl = st.file_uploader(
        "Zgjidh audio (MP3, WAV, M4A, OGG, FLAC)",
        type=["mp3", "wav", "m4a", "mp4", "ogg", "flac"],
        accept_multiple_files=True,
        key="transcribe_upload"
    )

    if st.button("▶️ Transkripto", type="primary", disabled=not upl, key="transcribe_btn"):
        tmpdir = pathlib.Path("tmp_audio_transcribe")
        tmpdir.mkdir(exist_ok=True)
        paths = []
        for up in upl:
            p = tmpdir / up.name
            with open(p, "wb") as f:
                f.write(up.getbuffer())
            paths.append(p)

        st.info("⏳ Duke transkriptuar...")
        try:
            out = transcribe_audio_files(
                input_paths=paths,
                out_dir=OUT_DIR,
                session_name=session_name_transcribe or None,
                subpath="Transkripte",
                save_txt=True,
                save_docx=save_docx,
                reuse_existing=reuse_existing,
                force=force,
                keep_wav=False,
                auto_session_if_blank=True,
                agent_map=agent_mapping,
            )
            st.success(f"✅ Transkriptimi u krye për {len(out.get('txt_paths', []))} file.")
            st.info(f"📁 Folder output: {out.get('out_folder')}")
            for t in out.get("txt_paths", []):
                st.markdown(f"- {t.as_posix()}")

            # Model usage
            log_path = pathlib.Path("C:/vicidial_agent/out_analysis/model_usage_global.json")
            if log_path.exists():
                data = json.loads(log_path.read_text(encoding="utf-8"))
                a, b, c = data.get("gpt4o_direct", 0), data.get("gpt4o_fallback_wav", 0), data.get("whisper_fallback", 0)
                st.info(f"📊 **Model Usage:** 4o-direct={a} | 4o-fallback-wav={b} | whisper-fallback={c}")
        except Exception as e:
            st.error(f"❌ Gabim gjatë transkriptimit: {e}")

# ======================== TAB 4: MATERIALE AI ========================
with tab4:
    import pathlib
    import json
    from datetime import datetime, time
    from core.materials_generator import (
        generate_objections_and_responses,
        generate_sales_script,
        generate_faq,
        generate_best_practices,
        export_to_docx,
        export_to_json,
        export_to_txt
    )
    from core.campaign_manager import get_all_campaigns, get_campaign_hints
    from core.db_vicidial import list_recordings, set_db_connection, get_current_db_key, _read_db_secrets
    from core.downloader_vicidial import download_recording
    from core.transcription_audio import transcribe_audio_files
    from core.config import OUT_DIR
    import urllib3
    urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

    st.markdown("### 🤖 Materiale AI - Gjenerues Automatik")
    st.caption("Gjenero materiale trajnimi dhe shitjeje bazuar në analizën e telefonatave reale.")

    # Material type selector
    material_type = st.selectbox(
        "Zgjidh llojin e materialit",
        [
            "🎯 Objeksione & Përgjigje Konsultative",
            "📝 Skript Shitjeje",
            "❓ FAQ (Pyetje të Shpeshta)",
            "⭐ Best Practices",
            "🎵 Shkarko Audio për Trajnim (Top/Bottom Performers)"
        ],
        help="Zgjidhni çfarë materiali dëshironi të gjeneroni",
        key="mat_type"
    )

    st.markdown("---")

    # Campaign/Project selector
    st.markdown("#### 1️⃣ Zgjedh Projektin/Fushatën")
    campaigns = get_all_campaigns()

    campaign_options = {"Asnjë (pa kontekst specifik)": None}
    for camp in campaigns:
        campaign_options[camp.get("name", "Unknown")] = camp.get("id")

    selected_campaign_name = st.selectbox(
        "Projekt/Fushatë",
        options=list(campaign_options.keys()),
        help="Zgjidhni projektin për të marrë kontekstin dhe dokumentet",
        key="mat_campaign"
    )
    selected_campaign_id = campaign_options[selected_campaign_name]

    if selected_campaign_id:
        campaign_hints = get_campaign_hints(selected_campaign_id)
        st.success(f"✅ Projekt i zgjedhur: {selected_campaign_name}")
        with st.expander("📄 Shiko kontekstin e projektit"):
            st.write(f"**Kontekst:** {campaign_hints.get('project_context_hint', 'Nuk ka')}")
            st.write(f"**Dokumente ({len(campaign_hints.get('documents_text', ''))} karaktere):** {'Po' if campaign_hints.get('documents_text') else 'Jo'}")
    else:
        campaign_hints = {"project_context_hint": "", "documents_text": "", "summary_hint": "", "bullets_hint": ""}

    st.markdown("---")

    # Source selection
    st.markdown("#### 2️⃣ Burimi i të Dhënave")

    source_type = st.radio(
        "Nga ku të merren transkriptet?",
        [
            "📁 Transkripte ekzistuese (folder lokal)",
            "🎙️ Regjistrime nga DB (download + transkriptim automatik)",
            "🎵 Ngarko audio lokal (upload + transkriptim automatik)"
        ],
        key="mat_source"
    )

    transcript_paths = []

    if "🎵 Ngarko audio" in source_type:
        # Upload audio files directly
        st.markdown("**Ngarko audio files**")

        uploaded_audio = st.file_uploader(
            "Zgjidh audio files (MP3, WAV, M4A, OGG, FLAC)",
            type=["mp3", "wav", "m4a", "mp4", "ogg", "flac"],
            accept_multiple_files=True,
            help="Mund të ngarkosh më shumë audio files njëherësh",
            key="mat_upload_audio"
        )

        if uploaded_audio:
            st.info(f"📎 U zgjodhën {len(uploaded_audio)} audio files")

            # Show file details
            with st.expander(f"📋 Shiko detajet e file-ave ({len(uploaded_audio)} files)"):
                total_size = 0
                for up in uploaded_audio:
                    size_mb = len(up.getvalue()) / (1024 * 1024)
                    total_size += size_mb
                    st.text(f"- {up.name} ({size_mb:.2f} MB)")
                st.caption(f"📊 Madhësia totale: {total_size:.2f} MB")

            if st.button("🎤 Transkripto Audio", type="primary", key="mat_transcribe_uploaded"):
                # Create temp directory
                temp_upload_dir = pathlib.Path("tmp_upload_materials")
                temp_upload_dir.mkdir(exist_ok=True)

                # Save uploaded files
                audio_paths = []
                prog = st.progress(0, text="Duke ruajtur audio files...")

                for i, up in enumerate(uploaded_audio):
                    out_path = temp_upload_dir / up.name
                    with open(out_path, "wb") as f:
                        f.write(up.getbuffer())
                    audio_paths.append(out_path)
                    prog.progress((i + 1) / len(uploaded_audio), text=f"Ruajtur: {i + 1}/{len(uploaded_audio)}")

                st.success(f"✅ U ruajtën {len(audio_paths)} audio files")

                # Transcribe
                if audio_paths:
                    st.info(f"⏳ Duke transkriptuar {len(audio_paths)} audio files...")

                    try:
                        temp_session = f"materials_upload_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

                        result = transcribe_audio_files(
                            input_paths=audio_paths,
                            out_dir=OUT_DIR,
                            session_name=temp_session,
                            subpath="Transkripte",
                            save_txt=True,
                            save_docx=False,
                            reuse_existing=True,
                            force=False,
                            keep_wav=False,
                            auto_session_if_blank=True
                        )

                        transcript_paths = result.get("txt_paths", [])
                        st.success(f"✅ U transkriptuan {len(transcript_paths)} audio files")
                        st.session_state['mat_transcript_paths'] = transcript_paths
                        st.info(f"📁 Transkriptet u ruajtën te: {result.get('out_folder')}")

                        # Show model usage
                        log_path = pathlib.Path("C:/vicidial_agent/out_analysis/model_usage_global.json")
                        if log_path.exists():
                            data = json.loads(log_path.read_text(encoding="utf-8"))
                            a, b, c = data.get("gpt4o_direct", 0), data.get("gpt4o_fallback_wav", 0), data.get("whisper_fallback", 0)
                            st.info(f"📊 **Model Usage:** 4o-direct={a} | 4o-fallback-wav={b} | whisper-fallback={c}")

                    except Exception as e:
                        st.error(f"❌ Gabim gjatë transkriptimit: {e}")
                        import traceback
                        st.code(traceback.format_exc())

    elif "📁 Transkripte" in source_type:
        # Existing transcripts from folder
        st.markdown("**Zgjidhni folderin me transkripte**")

        transcript_folder = st.text_input(
            "Path i folderit",
            value=str(OUT_DIR),
            help="Folderi që përmban transkriptet (.txt)",
            key="mat_transcript_folder"
        )

        include_subdirs = st.checkbox(
            "Përfshi edhe nënfolderat",
            value=True,
            help="Kërko transkripte në të gjitha nënfolderat",
            key="mat_subdirs"
        )

        if st.button("🔍 Gjej transkriptet", key="mat_find_transcripts"):
            folder = pathlib.Path(transcript_folder)
            if folder.exists():
                if include_subdirs:
                    transcript_paths = list(folder.rglob("*.txt"))
                else:
                    transcript_paths = list(folder.glob("*.txt"))

                # Filter out very small files
                transcript_paths = [p for p in transcript_paths if p.stat().st_size > 100]

                if transcript_paths:
                    st.success(f"✅ U gjetën {len(transcript_paths)} transkripte")
                    st.session_state['mat_transcript_paths'] = transcript_paths

                    # Show sample
                    with st.expander(f"📋 Shiko listën ({min(10, len(transcript_paths))} të parët)"):
                        for p in transcript_paths[:10]:
                            st.text(f"- {p.name} ({p.stat().st_size // 1024} KB)")
                else:
                    st.warning("⚠️ Nuk u gjend asnjë transkript në këtë folder")
            else:
                st.error("❌ Folderi nuk ekziston")

    elif "🎙️ Regjistrime" in source_type:
        # Download from DB + transcribe
        st.markdown("**Filtrat për regjistrimet**")

        # Database selector
        col_db1, col_db2 = st.columns([3, 1])
        with col_db1:
            db1_host, _, _, _ = _read_db_secrets("db")
            db2_host, _, _, _ = _read_db_secrets("db2")
            db_options_mat = {
                f"Database 1 ({db1_host or 'Default'})": "db",
                f"Database 2 ({db2_host or '95.217.87.125'})": "db2"
            }
            db_label_mat = st.selectbox(
                "Databaza",
                options=list(db_options_mat.keys()),
                index=0 if get_current_db_key() == "db" else 1,
                key="mat_db_select"
            )
            selected_db_key_mat = db_options_mat[db_label_mat]
            set_db_connection(selected_db_key_mat)
        with col_db2:
            current_host_mat, _, _, _ = _read_db_secrets(selected_db_key_mat)
            st.caption(f"🔌 **{current_host_mat}**")

        # Date/time filters
        col_d1, col_d2 = st.columns(2)
        with col_d1:
            mat_start_date = st.date_input("Data fillimit", key="mat_start_date")
            mat_start_time = st.time_input("Ora fillimit", value=time(0,0,0), key="mat_start_time")
        with col_d2:
            mat_end_date = st.date_input("Data mbarimit", key="mat_end_date")
            mat_end_time = st.time_input("Ora mbarimit", value=time(23,59,59), key="mat_end_time")

        mat_campaign_filter = st.text_input(
            "Fushata Vicidial (opsional)",
            help="Filtro vetëm regjistrime nga kjo fushatë",
            key="mat_campaign_filter"
        )

        # Duration filter
        col_dur1, col_dur2 = st.columns(2)
        with col_dur1:
            mat_min_dur = st.number_input("Kohëzgjatja min (sek)", 0, 3600, 60, key="mat_min_dur")
        with col_dur2:
            mat_max_dur = st.number_input("Kohëzgjatja max (sek)", 0, 3600, 0, help="0 = pa kufi", key="mat_max_dur")

        mat_max_files = st.number_input(
            "Maksimumi i regjistrimeve",
            10, 10000, 200,
            help="Sa regjistrime të shkarkosh dhe transkriptosh",
            key="mat_max_files"
        )

        # Auth
        col_a1, col_a2 = st.columns(2)
        with col_a1:
            mat_auth_user = st.text_input("Basic auth user", key="mat_auth_user")
        with col_a2:
            mat_auth_pass = st.text_input("Basic auth pass", type="password", key="mat_auth_pass")

        # Fshij butonin "Shkarko & Transkripto" për material type "Shkarko Audio për Trajnim"
        # Në këtë rast, operacioni do të bëhet direkt nga butoni kryesor
        if "🎵 Shkarko Audio" not in material_type:
            if st.button("⬇️ Shkarko & Transkripto", type="primary", key="mat_download_btn"):
                start_dt = datetime.combine(mat_start_date, mat_start_time).strftime("%Y-%m-%d %H:%M:%S")
                end_dt = datetime.combine(mat_end_date, mat_end_time).strftime("%Y-%m-%d %H:%M:%S")

                st.info(f"⏳ Duke kërkuar regjistrimet: {start_dt} → {end_dt}")

                try:
                    rows = list_recordings(
                        start_dt,
                        end_dt,
                        mat_campaign_filter.strip() or None,
                        limit=int(mat_max_files)
                    )
                except Exception as e:
                    st.error(f"❌ Gabim në DB: {e}")
                    st.stop()

                if not rows:
                    st.warning("⚠️ Nuk u gjend asnjë regjistrim")
                    st.stop()

                # Download
                temp_session = f"materials_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                temp_root = OUT_DIR / temp_session

                downloaded = 0
                filtered = 0
                audio_files = []

                prog = st.progress(0, text="Duke shkarkuar...")
                total = len(rows)

                for i, r in enumerate(rows, start=1):
                    filename = r.get("filename", "recording")
                    location = r.get("location", "")
                    length_sec = r.get("length_in_sec", 0)
                    user = (r.get("user") or "UNKNOWN").strip().title()
                    campaign_id = (r.get("campaign_id") or "UNKNOWN").strip()

                    # Duration filter
                    if mat_min_dur > 0 and length_sec < mat_min_dur:
                        filtered += 1
                        continue
                    if mat_max_dur > 0 and length_sec > mat_max_dur:
                        filtered += 1
                        continue

                    ext = pathlib.Path(location).suffix or ".wav"
                    safe_user = user.replace("/", "-")
                    safe_campaign = campaign_id.replace("/", "-")

                    out_dir = temp_root / safe_campaign / safe_user
                    out_file = out_dir / f"{filename}{ext}"

                    try:
                        auth = (mat_auth_user, mat_auth_pass) if mat_auth_user or mat_auth_pass else None
                        ok = download_recording(location, out_file, auth=auth)
                        if ok:
                            downloaded += 1
                            audio_files.append(out_file)
                    except Exception:
                        pass

                    prog.progress(int(i/total*100), text=f"Shkarkim: {i}/{total}")

                st.success(f"✅ U shkarkuan {downloaded} regjistrime ({filtered} u filtruan)")

                if audio_files:
                    st.info(f"⏳ Duke transkriptuar {len(audio_files)} audio...")

                    try:
                        result = transcribe_audio_files(
                            input_paths=audio_files,
                            out_dir=OUT_DIR,
                            session_name=temp_session,
                            subpath="Transkripte",
                            save_txt=True,
                            save_docx=False,
                            reuse_existing=True,
                            force=False,
                            keep_wav=False,
                            auto_session_if_blank=True
                        )

                        transcript_paths = result.get("txt_paths", [])
                        st.success(f"✅ U transkriptuan {len(transcript_paths)} audio")
                        st.session_state['mat_transcript_paths'] = transcript_paths

                    except Exception as e:
                        st.error(f"❌ Gabim në transkriptim: {e}")
                        st.stop()

    st.markdown("---")

    # Selection parameters (only for "Shkarko Audio për Trajnim")
    if "🎵 Shkarko Audio" in material_type:
        st.markdown("#### 3️⃣ Parametrat e Zgjedhjes")

        col1, col2 = st.columns(2)
        with col1:
            max_to_analyze = st.number_input(
                "Maksimumi i regjistrimeve për analizë",
                10, 10000, 500,
                help="Sa regjistrime të merren për analizë (të gjitha)",
                key="mat_max_analyze"
            )

        with col2:
            max_to_download = st.number_input(
                "Sa të perzgjidhen për shkarkim",
                5, 100, 20,
                help="Sa regjistrime të zgjedhura të shkarkohen",
                key="mat_max_download"
            )

        # Kriteri i zgjedhjes
        selection_criteria = st.selectbox(
            "Kriteri i zgjedhjes",
            [
                "⭐ Më të mirat (score >= 4.0) - për imitim",
                "⚠️ Më problematiket (score <= 2.5) - për korrigjim",
                "📊 Mesataret (score 3.0-4.0) - për analizë",
                "🎯 Të balancuara (50% top + 50% bottom)"
            ],
            key="mat_selection_criteria"
        )

        st.info("💡 Regjistrimet e zgjedhura do të ruhen në folder të organizuar bazuar në projektin e zgjedhur.")

        st.markdown("---")

        # Additional instructions
        st.markdown("#### 4️⃣ Instruksione Shtesë (Opsionale)")
    else:
        # Additional instructions
        st.markdown("#### 3️⃣ Instruksione Shtesë (Opsionale)")

    additional_instructions = st.text_area(
        "Instruksione specifike për gjenerimin e materialit",
        placeholder="P.sh. 'Fokuso tek objeksionet teknike', 'Përfshi shembuj konkretë', 'Ton profesional por miqësor', etj.",
        height=100,
        help="Këto instruksione do t'i jepen AI-së për të personalizuar output-in",
        key="mat_instructions"
    )

    # Language
    mat_language = st.selectbox(
        "Gjuha e materialit",
        ["🇦🇱 Shqip", "🇮🇹 Italisht", "🇬🇧 Anglisht"],
        key="mat_language"
    )

    lang_code = {"🇦🇱 Shqip": "sq", "🇮🇹 Italisht": "it", "🇬🇧 Anglisht": "en"}[mat_language]

    st.markdown("---")

    # Generate button
    if "🎵 Shkarko Audio" in material_type:
        st.markdown("#### 5️⃣ Gjenero & Shkarko Regjistrime")
    else:
        st.markdown("#### 4️⃣ Gjenero Materialin")

    # Check if we have transcripts
    # Për "Shkarko Audio" + "DB", nuk kemi nevojë për transkripte në session_state (do të merren direkt nga DB)
    if "🎵 Shkarko Audio" in material_type and "🎙️ Regjistrime" in source_type:
        has_transcripts = True  # Do të merren direkt nga DB
    else:
        has_transcripts = 'mat_transcript_paths' in st.session_state and st.session_state['mat_transcript_paths']

    if not has_transcripts:
        if "🎵 Shkarko Audio" in material_type:
            st.warning("⚠️ Nuk ka transkripte të ngarkuara. Përdor butonin 'Gjej transkriptet', 'Transkripto Audio', ose zgjidh 'Regjistrime nga DB' më sipër.")
        else:
            st.warning("⚠️ Nuk ka transkripte të ngarkuara. Përdor butonin 'Gjej transkriptet' ose 'Shkarko & Transkripto' më sipër.")

    # Button text changes based on material type
    if "🎵 Shkarko Audio" in material_type:
        button_text = "🚀 Analizo & Shkarko Regjistrime"
    else:
        button_text = "🚀 Gjenero Materialin"

    generate_btn = st.button(
        button_text,
        type="primary",
        disabled=not has_transcripts,
        key="mat_generate_btn"
    )

    if generate_btn and has_transcripts:
        # Merr transkriptet nga session_state (për burime të tjera)
        # Për "Shkarko Audio" + "DB", transcript_paths do të merret më poshtë
        if "🎵 Shkarko Audio" not in material_type or "🎙️ Regjistrime" not in source_type:
            transcript_paths = st.session_state['mat_transcript_paths']
            st.info(f"⏳ Duke gjeneruar materialin nga {len(transcript_paths)} transkripte...")

        # Prepare context
        context = campaign_hints.get("project_context_hint", "")
        documents = campaign_hints.get("documents_text", "")

        # Generate based on type
        try:
            if "Objeksione" in material_type:
                result = generate_objections_and_responses(
                    transcript_paths=transcript_paths,
                    campaign_context=context,
                    documents_text=documents,
                    additional_instructions=additional_instructions,
                    language=lang_code,
                    min_objections=10
                )
                export_type = "objections"

            elif "Skript" in material_type:
                result = generate_sales_script(
                    transcript_paths=transcript_paths,
                    campaign_context=context,
                    documents_text=documents,
                    additional_instructions=additional_instructions,
                    language=lang_code
                )
                export_type = "script"

            elif "FAQ" in material_type:
                result = generate_faq(
                    transcript_paths=transcript_paths,
                    campaign_context=context,
                    documents_text=documents,
                    additional_instructions=additional_instructions,
                    language=lang_code
                )
                export_type = "faq"

            elif "🎵 Shkarko Audio" in material_type:
                # Import the new function
                from core.materials_generator import select_recordings_for_training

                # Nëse burimi është DB, duhet të shkarkojmë dhe transkriptojmë fillimisht
                if "🎙️ Regjistrime" in source_type:
                    # HAPI 1/4: Shkarko nga DB
                    st.info("📥 **Hapi 1/4:** Duke shkarkuar regjistrimet nga DB...")

                    start_dt = datetime.combine(mat_start_date, mat_start_time).strftime("%Y-%m-%d %H:%M:%S")
                    end_dt = datetime.combine(mat_end_date, mat_end_time).strftime("%Y-%m-%d %H:%M:%S")

                    try:
                        rows = list_recordings(
                            start_dt,
                            end_dt,
                            mat_campaign_filter.strip() or None,
                            limit=int(mat_max_files)
                        )
                    except Exception as e:
                        st.error(f"❌ Gabim në DB: {e}")
                        st.stop()

                    if not rows:
                        st.warning("⚠️ Nuk u gjend asnjë regjistrim")
                        st.stop()

                    # Download
                    temp_session = f"materials_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                    temp_root = OUT_DIR / temp_session

                    downloaded = 0
                    filtered = 0
                    audio_files = []

                    prog1 = st.progress(0, text="Duke shkarkuar...")
                    total = len(rows)

                    for i, r in enumerate(rows, start=1):
                        filename = r.get("filename", "recording")
                        location = r.get("location", "")
                        length_sec = r.get("length_in_sec", 0)
                        user = (r.get("user") or "UNKNOWN").strip().title()
                        campaign_id = (r.get("campaign_id") or "UNKNOWN").strip()

                        # Duration filter
                        if mat_min_dur > 0 and length_sec < mat_min_dur:
                            filtered += 1
                            continue
                        if mat_max_dur > 0 and length_sec > mat_max_dur:
                            filtered += 1
                            continue

                        ext = pathlib.Path(location).suffix or ".wav"
                        safe_user = user.replace("/", "-")
                        safe_campaign = campaign_id.replace("/", "-")

                        out_dir = temp_root / safe_campaign / safe_user
                        out_file = out_dir / f"{filename}{ext}"

                        try:
                            auth = (mat_auth_user, mat_auth_pass) if mat_auth_user or mat_auth_pass else None
                            ok = download_recording(location, out_file, auth=auth)
                            if ok:
                                downloaded += 1
                                audio_files.append(out_file)
                        except Exception:
                            pass

                        prog1.progress(int(i/total*100), text=f"Shkarkim: {i}/{total}")

                    st.success(f"✅ U shkarkuan {downloaded} regjistrime ({filtered} u filtruan)")

                    # HAPI 2/4: Transkripto
                    if audio_files:
                        st.info("🎤 **Hapi 2/4:** Duke transkriptuar audio files...")

                        try:
                            trans_result = transcribe_audio_files(
                                input_paths=audio_files,
                                out_dir=OUT_DIR,
                                session_name=temp_session,
                                subpath="Transkripte",
                                save_txt=True,
                                save_docx=False,
                                reuse_existing=True,
                                force=False,
                                keep_wav=False,
                                auto_session_if_blank=True
                            )

                            transcript_paths = trans_result.get("txt_paths", [])
                            st.success(f"✅ U transkriptuan {len(transcript_paths)} audio files")

                        except Exception as e:
                            st.error(f"❌ Gabim në transkriptim: {e}")
                            st.stop()
                    else:
                        st.error("❌ Nuk u shkarkuan asnjë audio file")
                        st.stop()

                # HAPI 3/4: Analizo dhe zgjedh
                st.info("🔍 **Hapi 3/4:** Duke analizuar dhe duke zgjedhur regjistrimet...")

                # Determine selection criteria
                if "Më të mirat" in selection_criteria:
                    criteria = "best"
                elif "Më problematiket" in selection_criteria:
                    criteria = "worst"
                elif "Mesataret" in selection_criteria:
                    criteria = "average"
                else:  # Të balancuara
                    criteria = "balanced"

                result = select_recordings_for_training(
                    transcript_paths=transcript_paths[:max_to_analyze],
                    max_to_analyze=max_to_analyze,
                    max_to_download=max_to_download,
                    selection_criteria=criteria,
                    campaign_context=context,
                    documents_text=documents,
                    language=lang_code,
                    campaign_name=selected_campaign_name or "UNKNOWN"
                )

                # HAPI 4/4: Organizimi (bëhet automatikisht brenda funksionit)
                st.info("📁 **Hapi 4/4:** Duke organizuar regjistrimet në folder...")

                export_type = "training_recordings"

            else:  # Best Practices
                result = generate_best_practices(
                    transcript_paths=transcript_paths,
                    campaign_context=context,
                    documents_text=documents,
                    additional_instructions=additional_instructions,
                    language=lang_code
                )
                export_type = "best_practices"

            if "error" in result:
                st.error(f"❌ Gabim: {result['error']}")
            else:
                st.success(f"✅ Materiali u gjenerua me sukses!")

                # Store in session
                st.session_state['mat_result'] = result
                st.session_state['mat_export_type'] = export_type

                # Display results
                st.markdown("---")
                st.markdown("### 📊 Rezultati")

                with st.expander("🔍 Shiko të dhënat e plota (JSON)", expanded=False):
                    st.json(result)

                # Formatted display based on type
                if export_type == "objections":
                    st.markdown(f"#### Objeksione të Gjeneruara: {len(result.get('objections', []))}")

                    for i, obj in enumerate(result.get('objections', []), 1):
                        with st.expander(f"{i}. [{obj.get('category')}] {obj.get('objection')}", expanded=(i <= 3)):
                            st.markdown(f"**Frekuencë:** {obj.get('frequency', 'N/A')}")
                            st.markdown(f"**Kontekst:** {obj.get('context', 'N/A')}")

                            approach = obj.get('consultative_approach', {})
                            st.markdown("##### Qasja Konsultative:")
                            st.markdown(f"**Prevention:** {approach.get('prevention', 'N/A')}")
                            st.markdown(f"**Value Building:** {approach.get('value_building', 'N/A')}")
                            st.markdown(f"**Response Framework:** {approach.get('response_framework', 'N/A')}")
                            st.markdown(f"**Shembull Dialogu:**\n```\n{approach.get('example_dialogue', 'N/A')}\n```")

                    st.markdown("---")
                    st.markdown("#### Strategji e Përgjithshme")
                    st.info(result.get('general_strategy', 'N/A'))

                    st.markdown("#### Rekomandime për Call Flow")
                    st.info(result.get('call_flow_recommendations', 'N/A'))

                elif export_type == "script":
                    script = result.get('script', {})
                    st.markdown("#### Skripti i Gjeneruar")

                    for section_name, section_data in script.items():
                        with st.expander(f"📌 {section_name.upper()}", expanded=True):
                            if isinstance(section_data, dict):
                                for key, value in section_data.items():
                                    st.markdown(f"**{key.replace('_', ' ').title()}:**")
                                    if isinstance(value, list):
                                        for item in value:
                                            st.markdown(f"- {item}")
                                    else:
                                        st.write(value)
                            else:
                                st.write(section_data)

                elif export_type == "faq":
                    st.markdown(f"#### FAQ: {len(result.get('faqs', []))} pyetje")

                    for i, faq in enumerate(result.get('faqs', []), 1):
                        with st.expander(f"{i}. {faq.get('question', 'N/A')}", expanded=(i <= 5)):
                            st.markdown(f"**Kategoria:** {faq.get('category', 'N/A')}")
                            st.markdown(f"**Frekuencë:** {faq.get('frequency', 'N/A')}")
                            st.markdown(f"**Përgjigje:** {faq.get('answer', 'N/A')}")
                            st.markdown(f"**Detaje:** {faq.get('detailed_answer', 'N/A')}")

                elif export_type == "best_practices":
                    st.markdown(f"#### Best Practices: {len(result.get('best_practices', []))}")

                    for i, practice in enumerate(result.get('best_practices', []), 1):
                        with st.expander(f"{i}. [{practice.get('category')}] {practice.get('practice')}", expanded=(i <= 5)):
                            st.markdown(f"**Shpjegim:** {practice.get('explanation', 'N/A')}")
                            st.markdown(f"**Shembull:** {practice.get('example', 'N/A')}")
                            st.markdown(f"**Vështirësi:** {practice.get('difficulty', 'N/A')}")

                elif export_type == "training_recordings":
                    st.markdown(f"#### 🎵 Regjistrime të Zgjedhura: {len(result.get('selected_recordings', []))}")

                    # Auto-fallback message
                    if result.get('auto_fallback', False):
                        st.warning(result.get('auto_fallback_message', ''))

                    # Statistics
                    stats = result.get('statistics', {})
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.metric("Total të Analizuara", result.get('total_analyzed', 0))
                    with col2:
                        st.metric("Top Performers", stats.get('best_count', 0))
                    with col3:
                        st.metric("Bottom Performers", stats.get('worst_count', 0))

                    st.info(f"📁 **Folder:** `{result.get('download_folder', 'N/A')}`")
                    st.info(f"🎯 **Kriteri:** {result.get('selection_criteria', 'N/A')}")

                    # Show selected recordings
                    for i, rec in enumerate(result.get('selected_recordings', []), 1):
                        with st.expander(f"{i}. {rec.get('agent', 'UNKNOWN')} - Score: {rec.get('score', 0)} ⭐", expanded=(i <= 3)):
                            col_a, col_b = st.columns(2)
                            with col_a:
                                st.markdown(f"**Agjent:** {rec.get('agent', 'N/A')}")
                                st.markdown(f"**Score:** {rec.get('score', 0)}")
                                st.markdown(f"**Ranking:** #{rec.get('ranking', 'N/A')}")
                            with col_b:
                                st.markdown(f"**Audio:** `{rec.get('audio_path', 'N/A')}`")
                                st.markdown(f"**Transkript:** `{rec.get('transcript_path', 'N/A')}`")

                            st.markdown("**Arsyeja e zgjedhjes:**")
                            st.write(rec.get('reason', 'N/A'))

                            st.markdown("**Përmbledhje:**")
                            st.write(rec.get('summary', 'N/A'))

                            col_p1, col_p2 = st.columns(2)
                            with col_p1:
                                st.markdown("**Preggi:**")
                                for p in rec.get('preggi', []):
                                    st.markdown(f"- {p}")
                            with col_p2:
                                st.markdown("**Për të përmirësuar:**")
                                for p in rec.get('da_migliorare', []):
                                    st.markdown(f"- {p}")

                # Export options
                st.markdown("---")
                st.markdown("### 💾 Shkarko Materialin")
                st.caption("Kliko butonin për të shkarkuar materialin në kompjuterin tënd")

                col_exp1, col_exp2, col_exp3 = st.columns(3)

                with col_exp1:
                    # DOCX export
                    try:
                        from io import BytesIO
                        from docx import Document
                        from docx.shared import Pt, RGBColor
                        from docx.enum.text import WD_ALIGN_PARAGRAPH

                        doc = Document()

                        # Title
                        title = doc.add_heading(f"AI Generated: {export_type.upper()}", 0)
                        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

                        # Metadata
                        meta = doc.add_paragraph()
                        meta.add_run(f"Gjeneruar: {result.get('generated_at', 'N/A')}\n").italic = True
                        meta.add_run(f"Transkripte: {result.get('total_transcripts_analyzed', 0)}\n").italic = True
                        doc.add_paragraph("_" * 80)

                        # Content based on type
                        if export_type == "objections":
                            doc.add_heading("Objeksione dhe Përgjigje Konsultative", 1)
                            for obj in result.get("objections", []):
                                doc.add_heading(f"{obj.get('category', 'N/A')}: {obj.get('objection', 'N/A')}", 2)
                                doc.add_paragraph(f"Frekuencë: {obj.get('frequency', 'N/A')}")
                                doc.add_paragraph(f"Kontekst: {obj.get('context', 'N/A')}")
                                approach = obj.get("consultative_approach", {})
                                doc.add_heading("Qasja Konsultative:", 3)
                                doc.add_paragraph(f"Prevention: {approach.get('prevention', 'N/A')}")
                                doc.add_paragraph(f"Value Building: {approach.get('value_building', 'N/A')}")
                                doc.add_paragraph(f"Response Framework: {approach.get('response_framework', 'N/A')}")
                                doc.add_paragraph(f"Shembull:\n{approach.get('example_dialogue', 'N/A')}")
                                doc.add_paragraph()
                            doc.add_heading("Strategji e Përgjithshme", 1)
                            doc.add_paragraph(result.get("general_strategy", "N/A"))
                            doc.add_heading("Rekomandime për Flow", 1)
                            doc.add_paragraph(result.get("call_flow_recommendations", "N/A"))

                        elif export_type == "script":
                            script = result.get("script", {})
                            for section_name, section_data in script.items():
                                doc.add_heading(section_name.upper(), 1)
                                if isinstance(section_data, dict):
                                    for key, value in section_data.items():
                                        doc.add_heading(key.replace("_", " ").title(), 2)
                                        if isinstance(value, list):
                                            for item in value:
                                                doc.add_paragraph(str(item), style='List Bullet')
                                        else:
                                            doc.add_paragraph(str(value))

                        elif export_type == "faq":
                            for faq in result.get("faqs", []):
                                doc.add_heading(f"Q: {faq.get('question', 'N/A')}", 2)
                                doc.add_paragraph(f"Kategoria: {faq.get('category', 'N/A')}")
                                doc.add_paragraph(f"A: {faq.get('answer', 'N/A')}")
                                doc.add_paragraph(f"Detaje: {faq.get('detailed_answer', 'N/A')}")
                                doc.add_paragraph()

                        elif export_type == "best_practices":
                            for practice in result.get("best_practices", []):
                                doc.add_heading(f"{practice.get('category', 'N/A')}: {practice.get('practice', 'N/A')}", 2)
                                doc.add_paragraph(f"Shpjegim: {practice.get('explanation', 'N/A')}")
                                doc.add_paragraph(f"Shembull: {practice.get('example', 'N/A')}")
                                doc.add_paragraph(f"Vështirësi: {practice.get('difficulty', 'N/A')}")
                                doc.add_paragraph()

                        # Save to BytesIO
                        docx_buffer = BytesIO()
                        doc.save(docx_buffer)
                        docx_buffer.seek(0)

                        st.download_button(
                            label="📄 Shkarko DOCX",
                            data=docx_buffer,
                            file_name=f"{export_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="download_docx",
                            help="Kliko për të shkarkuar dokumentin Word"
                        )
                    except Exception as e:
                        st.error(f"❌ Gabim në DOCX: {e}")

                with col_exp2:
                    # JSON export
                    json_data = json.dumps(result, indent=2, ensure_ascii=False)
                    st.download_button(
                        label="📋 Shkarko JSON",
                        data=json_data,
                        file_name=f"{export_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                        mime="application/json",
                        key="download_json",
                        help="Kliko për të shkarkuar JSON"
                    )

                with col_exp3:
                    # TXT export
                    txt_data = f"=== AI GENERATED: {export_type.upper()} ===\n\n"
                    txt_data += f"Gjeneruar: {result.get('generated_at', 'N/A')}\n"
                    txt_data += f"Transkripte: {result.get('total_transcripts_analyzed', 0)}\n\n"
                    txt_data += "="*80 + "\n\n"
                    txt_data += json.dumps(result, indent=2, ensure_ascii=False)

                    st.download_button(
                        label="📝 Shkarko TXT",
                        data=txt_data,
                        file_name=f"{export_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                        mime="text/plain",
                        key="download_txt",
                        help="Kliko për të shkarkuar TXT"
                    )

        except Exception as e:
            st.error(f"❌ Gabim gjatë gjenerimit: {e}")
            import traceback
            st.code(traceback.format_exc())

st.markdown("---")
st.info("💡 **Shënim:** Për pipeline të plotë automatik, përdor faqen 'Home' ose '1_Analizë Automatike'.")

